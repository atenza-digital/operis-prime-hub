from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "cliente" / "ROTEIRO_VALIDACAO_COMERCIAL_TARCISIO.docx"
PROPOSAL_IMAGE = ROOT / "docs" / "evidencias" / "p0-propostas" / "proposta-ciperprag-padrao-v5-ritmo-page-1.png"
CONTRACT_IMAGE_1 = ROOT / "docs" / "evidencias" / "p0-contratos" / "contrato-ciperprag-padrao-v1-render-1.png"
CONTRACT_IMAGE_3 = ROOT / "docs" / "evidencias" / "p0-contratos" / "contrato-ciperprag-padrao-v1-render-3.png"

INK = RGBColor(11, 30, 45)
GREEN = RGBColor(11, 125, 83)
MUTED = RGBColor(90, 103, 125)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
            if row_index == 0:
                set_cell_shading(cell, "0B7D53")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.color.rgb = GREEN
    run.font.size = Pt(15 if level == 1 else 12)


def add_body(document: Document, text: str, *, indent: bool = True) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.12
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.75)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.left_indent = Cm(0.55)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_picture_if_exists(document: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        add_body(document, f"Imagem de referência não encontrada: {image_path}", indent=False)
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Cm(15.5))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(8)
    run = caption_paragraph.add_run(caption)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = MUTED


def add_key_value_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Informação"

    rows = [
        ("URL", "https://fieldops-homologacao.atenza.digital/login"),
        ("Usuário comercial", "homolog.comercial@atenza.digital"),
        ("Senha", "Usar a senha temporária informada separadamente ou solicitar reset antes do teste."),
        ("Ambiente", "Homologação. Não usar como produção."),
        ("Versão esperada", "0.6.3 ou superior em homologação."),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)
    style_table(table)


def add_flow_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Etapa"
    table.rows[0].cells[1].text = "Resultado esperado"

    rows = [
        ("Cliente", "Dados cadastrais completos e disponíveis para seleção na proposta."),
        ("Catálogo", "Serviços/produtos com unidade, quantidade, frequência, valores e escopo técnico."),
        ("Proposta", "Documento segue o padrão visual aprovado da Ciperprag, com logo e dados dinâmicos do tenant."),
        ("Aprovação da proposta", "Status da proposta muda corretamente e libera a geração da minuta."),
        ("Minuta", "Minuta nasce da proposta aprovada para revisão, negociação e aceite formal."),
        ("Contrato final", "Contrato nasce da minuta aprovada e libera a integração operacional."),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)
    style_table(table)


def add_checklist(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Item testado"
    table.rows[0].cells[1].text = "OK / Parcial / Erro"
    table.rows[0].cells[2].text = "Observações do teste"

    rows = [
        "Login com usuário comercial",
        "Listagem de clientes",
        "Cadastro/edição de cliente",
        "Catálogo de serviços/produtos",
        "Criação de proposta",
        "PDF da proposta",
        "Mudança de status da proposta",
        "Geração de minuta",
        "PDF da minuta",
        "Aprovação da minuta",
        "Geração do contrato final",
        "PDF do contrato final",
        "Integração do contrato com operacional",
        "Acentuação e datas no padrão brasileiro",
        "Fluxo lógico e sem burocracia",
    ]
    for item in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], item)
        set_cell_text(cells[1], "")
        set_cell_text(cells[2], "")
    style_table(table)


def add_response_area(document: Document) -> None:
    for label in [
        "O que funcionou bem:",
        "O que ficou confuso:",
        "Erros encontrados:",
        "Sugestões de melhoria:",
        "Prioridade percebida:",
    ]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(label)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(10.5)

        for _ in range(3):
            line = document.add_paragraph("_" * 95)
            line.paragraph_format.space_after = Pt(1)
            for run in line.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(190, 198, 210)


def build_document() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Roteiro de Validação Comercial")
    title_run.font.name = "Arial"
    title_run.font.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = INK

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Propostas, minutas e contratos - Atenza FieldOps")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = MUTED

    add_body(
        document,
        "Este roteiro deve ser usado para validar o fluxo comercial em homologação. "
        "As observações devem ser preenchidas neste próprio documento, com indicação clara "
        "do que ficou correto, parcial ou com erro. Evite enviar problemas soltos no grupo "
        "de WhatsApp, porque o documento consolidado facilita a análise e o retorno técnico.",
    )

    add_heading(document, "1. Acesso ao ambiente")
    add_key_value_table(document)

    add_heading(document, "2. Como registrar o resultado")
    for item in [
        "Marque cada item como OK, Parcial ou Erro.",
        "Quando houver erro, descreva o passo, o resultado esperado e o que aconteceu.",
        "Se possível, cole um print logo abaixo da observação.",
        "Ao final, envie este DOCX preenchido para a Atenza. Não concentrar as ocorrências no grupo de WhatsApp.",
    ]:
        add_bullet(document, item)

    add_heading(document, "3. Fluxo esperado do Comercial")
    add_body(
        document,
        "O fluxo correto é: cadastrar ou revisar cliente, cadastrar serviços ou produtos "
        "no catálogo, gerar proposta, enviar para o cliente, marcar a proposta como aprovada "
        "quando houver aceite comercial, gerar a minuta para revisão/negociação e, somente após "
        "a minuta aprovada, gerar o contrato final. Após o contrato ficar vigente, ele deve alimentar "
        "o operacional para agendamentos, consulta de saldo e geração de ordens de serviço.",
    )
    add_flow_table(document)

    add_heading(document, "4. Teste da proposta")
    for item in [
        "Acesse Comercial > Propostas e contratos.",
        "Crie uma proposta ou abra uma proposta existente.",
        "Confira cliente, serviços, frequência, quantidade, valor unitário, condições comerciais, validade e observações.",
        "Gere ou imprima o PDF e confira cabeçalho, rodapé, logo, dados da Ciperprag, dados do cliente, ordem das seções, acentuação, recuo dos parágrafos, texto justificado e paginação.",
        "Marque a proposta como enviada, em negociação, aprovada ou recusada conforme o caso de teste.",
    ]:
        add_bullet(document, item)
    add_picture_if_exists(document, PROPOSAL_IMAGE, "Referência visual da proposta para comparação durante o teste.")

    add_heading(document, "5. Teste da minuta e do contrato final")
    for item in [
        "A partir de uma proposta aprovada, clique em Gerar minuta.",
        "Confirme se os itens da proposta foram levados para a minuta sem redigitação desnecessária.",
        "Revise vigência, forma de pagamento, periodicidade, local, cláusulas, responsabilidades, observações e assinaturas.",
        "Aprove a minuta somente quando ela representar o acordo aceito pelo cliente.",
        "A partir da minuta aprovada, clique em Gerar contrato.",
        "Gere ou imprima o PDF do contrato final e confira se o documento segue o sistema visual aprovado, com cabeçalho, rodapé, seções em ordem e assinaturas alinhadas.",
        "Verifique se apenas o contrato final vigente fica disponível para o fluxo operacional.",
    ]:
        add_bullet(document, item)
    add_picture_if_exists(document, CONTRACT_IMAGE_1, "Referência visual da minuta/contrato - primeira página.")
    add_picture_if_exists(document, CONTRACT_IMAGE_3, "Referência visual da minuta/contrato - página de encerramento e assinaturas.")

    add_heading(document, "6. Checklist de validação")
    add_checklist(document)

    add_heading(document, "7. Campos para relato")
    add_response_area(document)

    add_heading(document, "8. Observações importantes")
    add_body(
        document,
        "A validação deve considerar que o sistema é SaaS. Logo, cores, logo, dados da empresa, "
        "responsáveis e documentos devem funcionar por tenant/configuração. Para a Ciperprag, "
        "os documentos devem seguir o padrão visual aprovado com os dados e identidade dela; "
        "para outros clientes, esses dados deverão ser parametrizados futuramente sem transformar "
        "o sistema em algo exclusivo da Ciperprag.",
    )
    add_body(
        document,
        "Qualquer divergência deve ser registrada neste arquivo. O grupo de WhatsApp deve ser usado "
        "apenas para avisar que o teste foi iniciado ou finalizado, não para concentrar prints e problemas soltos.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
