# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(
    r"C:\Projetos\Documentações\02_Templates\01_Operacionais\Template_Roteiro_Testes_Validacao_Atenza.docx"
)
OUT_DIR = ROOT / "docs" / "cliente" / "homologacao_roteiros"
QA_DIR = ROOT / "docs" / "evidencias" / "etapa7_homologacao" / "roteiro_final_tarcisio"
PRINTS_DIR = ROOT / "docs" / "evidencias" / "etapa7_homologacao" / "prints_roteiros"

DOCX_OUT = OUT_DIR / "Roteiro_Validacao_Final_Atenza_FieldOps_Tarcisio_v1.2.docx"
WHATSAPP_OUT = OUT_DIR / "MENSAGEM_WHATSAPP_HOMOLOGACAO_TARCISIO_v1.2.md"

VERSION = "0.6.3"
COMMIT = "5df8ccc"
URL = "https://fieldops-homologacao.atenza.digital/login"
DATE_BR = datetime.now().strftime("%d/%m/%Y")
PROJECT = "Atenza FieldOps"
CLIENT = "Ciperprag"
RESP_ATENZA = "Ewerton Gomes Almeida"
TESTER = "Tarcísio Lucas"

PROFILES = [
    ("Comercial", "homolog.comercial@atenza.digital", "Propostas, clientes, serviços, minutas e contratos"),
    ("Operacional", "homolog.operacao@atenza.digital", "Agenda, equipe, veículos, OS e encerramento"),
    ("Qualidade", "homolog.qualidade@atenza.digital", "Certificados, QR Code, histórico e validação pública"),
    ("Medição", "homolog.medicao@atenza.digital", "Medição, NF enviada, cobrança, pagamento e baixa manual no ERP"),
]

REFERENCE_DOCS = [
    (
        "Proposta",
        "docs/evidencias/qa_fluxo_visual/proposta-ciperprag.pdf",
        "Visual, cabeçalho/rodapé, parágrafos, condições comerciais, itens e paginação.",
    ),
    (
        "Contrato",
        "docs/evidencias/qa_fluxo_visual/contrato-ciperprag.pdf",
        "Assinaturas, cláusulas, dados das partes, serviços contratados e uso como base operacional.",
    ),
    (
        "Ordem de Serviço",
        "docs/evidencias/qa_fluxo_visual/os-atual.pdf",
        "Duas páginas quando possível, equipe, veículo, checklist, campos dinâmicos e legibilidade.",
    ),
    (
        "Certificado",
        "docs/cliente/certificados_montserrat/certificado-ciperprag-amostra-final.pdf",
        "Logo, marca d'água, fotos da OS, QR Code, hash, validade, assinatura e consistência da OS.",
    ),
    (
        "Relatório técnico",
        "docs/cliente/relatorios_tecnicos/relatorio-tecnico-ciperprag-amostra.pdf",
        "Dados da OS, checklist, fotos, produtos, EPIs, normas e ausência de valores comerciais.",
    ),
    (
        "Medição",
        "docs/evidencias/etapa7_homologacao/medicoes/med-validacao-2026-a4-retrato.pdf",
        "Consolidação por período, total, unidades, assinatura, rodapé e status financeiro-operacional.",
    ),
    (
        "Auditoria de fontes",
        "docs/evidencias/etapa7_homologacao/auditoria-fontes-documentais.md",
        "Confirma que os PDFs representativos usam Montserrat incorporada, sem fallback.",
    ),
]

PROFILE_STEPS = {
    "Comercial": [
        ("COM-01", "Entrar no sistema com o perfil Comercial.", "Acesso deve abrir as telas comerciais sem expor rotinas técnicas desnecessárias."),
        ("COM-02", "Conferir clientes e catálogo de serviços/produtos.", "Cliente e serviços devem estar fáceis de localizar e com textos legíveis."),
        ("COM-03", "Gerar uma proposta com itens, periodicidade, valores e condições.", "Proposta deve salvar e ficar disponível para visualização/impressão."),
        ("COM-04", "Aprovar a proposta e gerar minuta/contrato conforme fluxo disponível.", "Contrato deve liberar saldo operacional para agenda, sem retrabalho."),
        ("COM-05", "Abrir os PDFs de proposta e contrato de referência.", "Visual deve estar coerente com o padrão Ciperprag aprovado e sem fonte incorreta."),
    ],
    "Operacional": [
        ("OP-01", "Abrir Agenda/Agendamentos e selecionar contrato com saldo.", "Operacional não deve ver valores comerciais sensíveis."),
        ("OP-02", "Preencher data, equipe, veículo, local e tag/equipamento quando aplicável.", "Dados devem aparecer corretamente no agendamento e na OS."),
        ("OP-03", "Gerar OS a partir do agendamento.", "OS deve receber numeração automática e permitir impressão para campo."),
        ("OP-04", "Encerrar OS com execução, observações, tag/equipamento e até 3 fotos.", "OS encerrada deve alimentar certificado, histórico, relatório e medição."),
        ("OP-05", "Validar recorrência quando o serviço tiver frequência definida.", "Sistema deve sugerir próximo agendamento sem criar burocracia."),
    ],
    "Qualidade": [
        ("QL-01", "Abrir Certificados e Histórico.", "Tela deve permitir localizar OS encerrada e serviços com/sem certificado."),
        ("QL-02", "Gerar ou abrir certificado de OS elegível.", "Dados do certificado devem vir da mesma OS, tenant e cliente."),
        ("QL-03", "Conferir QR Code e rota pública de validação.", "Consulta pública deve exibir o mesmo certificado, hash e status."),
        ("QL-04", "Verificar fotos, produtos, licenças e assinatura.", "Blocos opcionais só devem aparecer quando houver dado/configuração."),
        ("QL-05", "Abrir relatório técnico quando aplicável.", "Relatório deve usar dados operacionais, sem valores comerciais."),
    ],
    "Medição": [
        ("MED-01", "Abrir Medição e filtrar cliente/período.", "Somente OS encerradas e elegíveis devem entrar na medição."),
        ("MED-02", "Gerar medição do período.", "Itens, unidades, quantidades e total devem ser coerentes."),
        ("MED-03", "Abrir PDF da medição.", "PDF deve estar em A4 retrato, com logo, cliente, itens, total e assinatura."),
        ("MED-04", "Atualizar status de NF enviada, aguardando pagamento e pago/baixado no ERP.", "Tela deve funcionar como acompanhamento, sem virar contas a receber."),
        ("MED-05", "Registrar se a informação ficou clara para a rotina administrativa.", "Fluxo deve ser simples para acompanhamento de cobrança e baixa manual."),
    ],
}


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "AtenzaNorticaBold" if bold else "AtenzaNortica"
    run.font.size = Pt(8.5)


def paragraph(doc: Document, text: str = "", *, justify: bool = True, indent: bool = True, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0.7) if indent else Cm(0)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "AtenzaNorticaBold" if bold else "AtenzaNortica"
    run.font.size = Pt(10)
    return p


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = text
    run.font.name = "AtenzaNeuePowerUltra" if level == 1 else "AtenzaNorticaBold"
    run.font.size = Pt(18 if level == 1 else 13)
    run.bold = True
    return p


def table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        set_cell_text(cell, header, True)
        shade(cell, "F2F2ED")
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths_cm:
        for row in tbl.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return tbl


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "AtenzaNortica"
    run.font.size = Pt(9.5)


def add_image_if_exists(doc: Document, title: str, relative_path: str, width: float = 5.9) -> None:
    path = ROOT / relative_path
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    r = cap.add_run(title)
    r.italic = True
    r.font.name = "AtenzaNortica"
    r.font.size = Pt(8)


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    clear_body(doc)

    paragraph(doc, "DOCUMENTO INSTITUCIONAL", justify=False, indent=False, bold=True)
    heading(doc, "Roteiro de Validação Final de Homologação", 1)
    paragraph(
        doc,
        "Atenza FieldOps - validação assistida do fluxo P0 da Ciperprag, cobrindo Comercial, Operacional, Qualidade e Medição.",
        justify=False,
        indent=False,
    )

    heading(doc, "Ficha do projeto", 2)
    table(
        doc,
        ["Campo", "Informação"],
        [
            ["Projeto", PROJECT],
            ["Cliente final", CLIENT],
            ["Responsável Atenza", RESP_ATENZA],
            ["Validador principal", TESTER],
            ["Ambiente", "Homologação"],
            ["URL", URL],
            ["Versão esperada", VERSION],
            ["Commit/evidência técnica", COMMIT],
            ["Data da rodada", DATE_BR],
            ["Status", "Rodada assistida para aceite interno"],
            ["Critério de aprovação", "Fluxo compreensível, documentos corretos e divergências registradas no próprio roteiro."],
        ],
        [4.2, 11.7],
    )

    heading(doc, "1. Como usar este roteiro", 2)
    paragraph(
        doc,
        "Execute os passos por perfil e marque cada item como Aprovado, Aprovado com ressalva ou Reprovado. Quando encontrar problema, registre o passo, o número do documento envolvido e cole um print diretamente neste arquivo. Evite enviar ocorrências soltas no grupo, porque o documento consolidado facilita a priorização das correções.",
    )
    bullet(doc, "Use dados fictícios quando precisar preencher observações, NF, anexos ou campos de teste.")
    bullet(doc, "Não informe senhas, tokens, chaves privadas ou dados sensíveis no roteiro.")
    bullet(doc, "Se uma tela estiver confusa mesmo funcionando, marque como Aprovado com ressalva.")
    bullet(doc, "Se o fluxo travar ou gerar documento/dado incorreto, marque como Reprovado.")

    heading(doc, "2. Acessos de homologação", 2)
    table(
        doc,
        ["Perfil", "Usuário", "Senha", "Escopo principal"],
        [[profile, user, "Enviada por canal seguro", scope] for profile, user, scope in PROFILES],
        [3.1, 4.7, 3.4, 5.3],
    )

    heading(doc, "3. Documentos de referência desta rodada", 2)
    paragraph(
        doc,
        "Os arquivos abaixo foram regenerados ou revalidados na etapa atual. Eles servem como referência para comparação visual e técnica durante a validação.",
    )
    table(
        doc,
        ["Documento", "Arquivo", "O que validar"],
        [[name, file_path, check] for name, file_path, check in REFERENCE_DOCS],
        [3.0, 6.2, 6.7],
    )

    heading(doc, "4. Roteiro por perfil", 2)
    for profile, _user, scope in PROFILES:
        heading(doc, f"4.{PROFILES.index((profile, _user, scope)) + 1}. Perfil {profile}", 3)
        paragraph(doc, scope, justify=False, indent=False)
        table(
            doc,
            ["ID", "Ação", "Resultado esperado", "Status", "Observação"],
            [[step_id, action, expected, "Pendente", ""] for step_id, action, expected in PROFILE_STEPS[profile]],
            [2.1, 5.2, 5.4, 2.2, 2.9],
        )

    heading(doc, "5. Situação técnica antes da rodada assistida", 2)
    paragraph(
        doc,
        "A validação automatizada e o smoke test de homologação foram executados pelo CI/CD no workflow de deploy. Os itens abaixo não precisam ser repetidos manualmente, mas devem ser considerados como evidência de que a base técnica está pronta para o aceite do usuário.",
    )
    table(
        doc,
        ["Verificação", "Resultado", "Evidência"],
        [
            ["Deploy em homologação", "Aprovado", "Workflow CI/CD 29756006104"],
            ["Auditoria E2E e isolamento SaaS", "Aprovado", "audit:e2e e saas:tri-tenant"],
            ["Integridade de anexos", "Aprovado", "Hashes reparados e download autenticado validado"],
            ["Numeração documental", "Aprovado", "Duplicidades corrigidas e auditoria executada"],
            ["Dados com acentuação corrompida", "Aprovado", "auditoria de dados sem ocorrências"],
            ["Smoke público", "Aprovado", "Rotas públicas e validação de certificado"],
        ],
        [4.8, 3.0, 8.1],
    )

    heading(doc, "6. Cinco validações manuais restantes", 2)
    paragraph(
        doc,
        "Esta é a última rodada manual da Etapa 7. Execute os cinco itens abaixo com usuário humano, sem pular etapas. Para cada item, registre Aprovado, Aprovado com ressalva ou Reprovado e anexe o print no próprio DOCX quando houver qualquer dúvida.",
    )
    table(
        doc,
        ["ID", "Como executar", "Resultado esperado", "Registro"],
        [
            ["FINAL-01", "Percorrer o fluxo completo: proposta, minuta, contrato final, contrato operacional, agenda, OS, encerramento, certificado, medição e recorrência.", "O fluxo avança sem retrabalho indevido; cada etapa reaproveita os dados corretos da etapa anterior.", "Status: ____\nObservação: ____"],
            ["FINAL-02", "Conferir datas, horas, moeda, acentos e textos em todas as telas e PDFs usados no fluxo.", "Formato brasileiro consistente: dd/mm/aaaa, hora local, valores em R$ e acentuação legível.", "Status: ____\nPrint: ____"],
            ["FINAL-03", "Pedir a um usuário que execute o fluxo sem orientação externa e observar cliques, mensagens, estados vazios e confirmações.", "A próxima ação fica clara; mensagens explicam o que ocorreu; não há telas sobrepostas ou etapas confusas.", "Status: ____\nObservação: ____"],
            ["FINAL-04", "Na OS, selecionar tags/equipamentos cadastrados e verificar o uso do dado no encerramento, certificado, relatório e histórico.", "A tag/ativo permanece coerente em todos os documentos e não gera texto cortado ou campo ilegível.", "Status: ____\nOS: ____"],
            ["FINAL-05", "Abrir um certificado em tela e impresso em A4; ler o QR Code com dois celulares e conferir a rota pública.", "Cada QR Code abre exatamente o certificado correspondente e mostra tenant, OS, cliente, serviço, execução, validade e hash.", "Status: ____\nCert.: ____"],
        ],
        [1.4, 6.2, 5.7, 2.7],
    )
    paragraph(
        doc,
        "Critério de encerramento: a Etapa 7 só pode ser considerada aprovada depois que os cinco itens estiverem preenchidos. Se houver ressalva, descreva o impacto e a evidência; a Atenza consolidará a decisão no roadmap sem perder o item.",
    )

    heading(doc, "7. Checklist final da rodada", 2)
    table(
        doc,
        ["Item", "Critério", "Status"],
        [
            ["Fluxo comercial", "Proposta aprovada gera minuta/contrato e libera saldo operacional.", "Pendente"],
            ["Fluxo operacional", "Agenda gera OS, OS encerra com evidências e alimenta documentos.", "Pendente"],
            ["Certificados", "QR Code, hash, fotos e dados da OS estão consistentes.", "Pendente"],
            ["Relatórios técnicos", "Relatório usa dados da OS e não expõe valores comerciais.", "Pendente"],
            ["Medição", "Consolida OS por período e acompanha NF/pagamento/ERP.", "Pendente"],
            ["Acentuação", "Telas e documentos não apresentam caracteres quebrados.", "Pendente"],
            ["Usabilidade", "Usuário consegue entender a próxima ação sem orientação externa.", "Pendente"],
        ],
        [3.8, 9.0, 3.0],
    )

    heading(doc, "8. Evidências para conferência", 2)
    paragraph(
        doc,
        "Use a tabela de documentos de referência da seção 3 para abrir os PDFs aprovados diretamente no computador. Evitei inserir miniaturas dos documentos neste roteiro porque elas ficam pequenas e podem induzir avaliação incorreta; a validação visual deve ser feita nos próprios PDFs.",
    )
    table(
        doc,
        ["Validação", "Como conferir"],
        [
            ["Interface", "Acessar a URL de homologação e conferir diretamente as telas com o perfil indicado."],
            ["Documentos", "Abrir os PDFs listados na seção 3 em tamanho real, conferindo logo, fonte, acentuação, quebras, rodapé e dados dinâmicos."],
            ["Divergências", "Registrar qualquer problema na seção 7 deste roteiro, com print colado no próprio DOCX."],
        ],
        [4.0, 11.8],
    )

    heading(doc, "9. Registro de divergências", 2)
    paragraph(
        doc,
        "Preencha uma linha para cada divergência encontrada. Se não houver problema, registre 'Sem divergências' na primeira linha.",
    )
    table(
        doc,
        ["ID", "Perfil", "Tela/documento", "O que aconteceu", "Número/print", "Severidade", "Status"],
        [
            ["HML-001", "", "", "", "", "", "Aberto"],
            ["HML-002", "", "", "", "", "", "Aberto"],
            ["HML-003", "", "", "", "", "", "Aberto"],
            ["HML-004", "", "", "", "", "", "Aberto"],
            ["HML-005", "", "", "", "", "", "Aberto"],
        ],
        [1.7, 2.2, 3.0, 4.5, 2.6, 2.1, 1.9],
    )

    heading(doc, "10. Resultado final", 2)
    table(
        doc,
        ["Campo", "Preenchimento"],
        [
            ["Resultado geral", "Aprovado / Aprovado com ressalva / Reprovado"],
            ["Nome de quem testou", ""],
            ["Data do teste", ""],
            ["Resumo das principais observações", ""],
            ["Pode seguir para próxima etapa?", "Sim / Não"],
        ],
        [5.0, 10.8],
    )

    heading(doc, "Histórico de versões", 2)
    table(
        doc,
        ["Versão", "Data", "Responsável", "Descrição"],
        [
            ["v1.2", DATE_BR, RESP_ATENZA, "Consolidação das evidências CI/CD e dos cinco testes manuais restantes para aceite da Etapa 7."],
        ],
        [2.0, 2.8, 4.2, 7.0],
    )

    doc.save(DOCX_OUT)

    message = """Pessoal, segue o roteiro v1.2 para a validação final da homologação do Atenza FieldOps.

Por favor, usem o DOCX para registrar tudo: OK, ressalvas, erros, prints e números de proposta/contrato/OS/certificado/medição. Evitem mandar ocorrências soltas no grupo, porque fica mais difícil consolidar e priorizar.

A senha deve ser usada apenas pelo canal seguro combinado. O CI/CD já validou a base técnica; concentrem-se nos cinco testes finais indicados no DOCX. Ao finalizar, devolvam o próprio documento preenchido."""
    WHATSAPP_OUT.write_text(message, encoding="utf-8")
    print(DOCX_OUT)
    print(WHATSAPP_OUT)


if __name__ == "__main__":
    build()
