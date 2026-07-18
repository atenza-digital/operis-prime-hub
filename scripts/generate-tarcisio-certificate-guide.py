from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "cliente" / "ROTEIRO_VALIDACAO_CERTIFICADOS_TARCISIO.docx"
SAMPLE_PDF = ROOT / "docs" / "cliente" / "certificados_montserrat" / "certificado-ciperprag-amostra-final.pdf"
SAMPLE_IMAGE = ROOT / "docs" / "cliente" / "certificados_montserrat" / "certificado-ciperprag-amostra-final-render-1.png"
COMBINED_PDF = ROOT / "docs" / "cliente" / "certificados_montserrat" / "certificados-montserrat-validacao-final.pdf"
LOGO_ATENZA = Path(r"C:\Projetos\Documentações\04_Assets\atenza_logo_dark_template.png")
SELO_ATENZA = Path(r"C:\Projetos\Documentações\04_Assets\atenza_tecnologia_com_proposito.png")

INK = RGBColor(3, 4, 9)
GREEN = RGBColor(0, 127, 102)
ACCENT = RGBColor(0, 255, 204)
BLUE = RGBColor(56, 80, 160)
MUTED = RGBColor(78, 90, 108)
LIGHT = "F4FBF8"
HEADER_FILL = "030409"
FONT_TITLE = "Neue Power Ultra"
FONT_BODY = "Nortica Typeface"
FONT_BODY_BOLD = "Nortica Typeface Bold"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=90, start=130, bottom=90, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold: bool = False, color: RGBColor | None = None, size: float = 9.2) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.08
    run = paragraph.add_run(text)
    run.font.name = FONT_BODY_BOLD if bold else FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    set_cell_margin(cell)


def style_table(table, *, header_fill: str = HEADER_FILL) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margin(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    run.font.name = FONT_BODY
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
                    run.font.size = Pt(9.2)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
                        run.font.name = FONT_BODY_BOLD


def add_run(paragraph, text: str, *, bold: bool = False, size: float = 10.4, color: RGBColor | None = None):
    run = paragraph.add_run(text)
    run.font.name = FONT_BODY_BOLD if bold else FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def add_paragraph(document: Document, text: str, *, indent: bool = True, bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.18
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.7)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(paragraph, bold_prefix, bold=True)
        add_run(paragraph, text[len(bold_prefix):])
    else:
        add_run(paragraph, text)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = FONT_TITLE if level == 1 else FONT_BODY_BOLD
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.bold = True
    run.font.color.rgb = GREEN if level == 1 else BLUE
    run.font.size = Pt(15 if level == 1 else 12)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.left_indent = Cm(0.65)
    paragraph.paragraph_format.line_spacing = 1.16
    add_run(paragraph, text, size=10.2)


def add_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Cm(17.0))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    left, right = header_table.rows[0].cells
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
    if LOGO_ATENZA.exists():
        left.paragraphs[0].add_run().add_picture(str(LOGO_ATENZA), width=Cm(3.2))
    else:
        add_run(left.paragraphs[0], "Atenza", bold=True, size=12)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if SELO_ATENZA.exists():
        right.paragraphs[0].add_run().add_picture(str(SELO_ATENZA), width=Cm(3.4))
    else:
        add_run(right.paragraphs[0], "Tecnologia com Propósito", bold=True, size=9)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(paragraph, "Atenza • CNPJ 56.067.166/0001-60 • atenza.digital • Roteiro de homologação", size=8.2, color=MUTED)


def add_info_table(document: Document) -> None:
    rows = [
        ("Projeto", "Atenza FieldOps"),
        ("Cliente de homologação", "Ciperprag Serviços"),
        ("Etapa validada", "P0.5 - Certificados, QR Code, hash, fotos e validação pública"),
        ("Ambiente", "Homologação"),
        ("URL", "https://fieldops-homologacao.atenza.digital/login"),
        ("Perfil recomendado", "Qualidade / responsável técnico"),
        ("Usuário de teste", "homolog.qualidade@atenza.digital"),
        ("Senha", "Enviada por canal seguro ou redefinida antes do teste"),
        ("PDF de referência", SAMPLE_PDF.name),
        ("Pacote de cenários", COMBINED_PDF.name),
        ("Critério de aprovação", "Sem divergência visual crítica, dados consistentes, QR Code legível e validação pública coerente."),
    ]
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Campo"
    table.rows[0].cells[1].text = "Informação"
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)
    style_table(table)


def add_result_table(document: Document) -> None:
    checks = [
        ("Acesso", "Entrar na URL de homologação com o usuário de Qualidade."),
        ("Listagem", "Acessar Certificados e Histórico e localizar certificados ou OS com certificado pendente."),
        ("Amostra final", "Abrir o PDF de referência enviado junto deste roteiro."),
        ("Layout", "Confirmar A4 paisagem, Montserrat, marca d'água lateral, logo dinâmica, assinatura e rodapé institucional."),
        ("Faixa de rastreabilidade", "Confirmar: Certificado nº 7303/2026 • OS nº 2677 • Execução: 22/06/2026."),
        ("Sem redundância", "Confirmar que não aparece CERT. CERT, Certificado Certificado ou OS OS."),
        ("Quadro superior direito", "Confirmar validade, QR Code, código curto e fingerprint SHA-256 abreviado."),
        ("Fingerprint", "Confirmar formato: SHA-256: 08E597799288…ED0B52DA."),
        ("Validação pública", "Abrir a rota do QR Code e conferir dados do certificado."),
        ("Hash completo", "Na validação pública, confirmar que aparece o SHA-256 completo persistido."),
        ("Dados dinâmicos", "Conferir tenant emissor, cliente, CNPJ, endereço, serviço, OS, tag, datas e responsável técnico."),
        ("Produtos", "Confirmar que produtos aparecem somente quando aplicáveis."),
        ("Fotos", "Confirmar que o certificado aceita zero a três fotos sem deformar ou cortar indevidamente."),
        ("Licenças", "Confirmar que licenças aparecem somente quando configuradas."),
        ("Validade", "Confirmar que validade aparece somente quando aplicável."),
        ("Acentuação", "Conferir certificado, validação, execução, responsáveis, órgãos, químico, serviço e endereço."),
        ("Impressão", "Imprimir em A4 paisagem e conferir legibilidade geral."),
        ("QR impresso - aparelho 1", "Ler o QR Code impresso com um celular e confirmar que abre o certificado correto."),
        ("QR impresso - aparelho 2", "Ler o QR Code impresso com outro celular e confirmar o mesmo resultado."),
        ("QR em tela", "Ler o QR Code diretamente no PDF em tela e confirmar o certificado correspondente."),
        ("Histórico", "Confirmar que o histórico mostra serviços com certificado e sem certificado, quando aplicável."),
    ]
    table = document.add_table(rows=1, cols=4)
    headers = ["Item", "Resultado esperado", "OK / Parcial / Erro", "Observações / print"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for item, expected in checks:
        cells = table.add_row().cells
        set_cell_text(cells[0], item, bold=True, size=8.8)
        set_cell_text(cells[1], expected, size=8.8)
        set_cell_text(cells[2], "", size=8.8)
        set_cell_text(cells[3], "", size=8.8)
    style_table(table)


def add_qr_validation_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=3)
    headers = ["Leitura", "O que deve acontecer", "Resultado / observação"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    rows = [
        ("PDF em tela", "QR abre a página pública do certificado exato.", ""),
        ("PDF impresso - aparelho 1", "QR abre a mesma página pública, sem erro ou certificado trocado.", ""),
        ("PDF impresso - aparelho 2", "QR abre a mesma página pública, confirmando leitura por outro dispositivo.", ""),
        ("Comparação da página pública", "Número, OS, cliente, serviço, execução, validade e SHA-256 batem com o PDF.", ""),
    ]
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, bold=index == 0)
    style_table(table)


def add_response_area(document: Document) -> None:
    labels = [
        "O que ficou correto:",
        "O que ficou parcial:",
        "Erros encontrados:",
        "Sugestões de melhoria:",
        "Prints ou links de evidência:",
        "Parecer final do teste:",
    ]
    for label in labels:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(3)
        add_run(paragraph, label, bold=True, size=10.4, color=INK)
        for _ in range(3):
            line = document.add_paragraph("_" * 94)
            line.paragraph_format.space_after = Pt(1)
            for run in line.runs:
                run.font.name = FONT_BODY
                run.font.size = Pt(8.8)
                run.font.color.rgb = RGBColor(185, 194, 206)


def add_sample_image(document: Document) -> None:
    add_heading(document, "5. Amostra visual aprovada para comparação")
    add_paragraph(
        document,
        "Use a imagem abaixo apenas como referência visual rápida. A validação oficial deve ser feita no PDF anexado, pois nele é possível conferir QR Code, seleção de texto, impressão e metadados.",
    )
    if SAMPLE_IMAGE.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(SAMPLE_IMAGE), width=Cm(16.2))
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(caption, "Amostra final do certificado Ciperprag para validação visual.", size=8.4, color=MUTED)
    else:
        add_paragraph(document, f"Imagem não encontrada: {SAMPLE_IMAGE}", indent=False)


def build_document() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)

    styles = document.styles
    styles["Normal"].font.name = FONT_BODY
    styles["Normal"].font.size = Pt(10.4)

    add_header_footer(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Roteiro de Validação de Certificados")
    run.font.name = FONT_TITLE
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = INK

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    add_run(subtitle, "Atenza FieldOps - Homologação Ciperprag", size=11.2, color=MUTED)

    add_paragraph(
        document,
        "Este roteiro deve ser usado para validar a etapa de certificados em homologação. O objetivo é confirmar se o certificado final aprovado mantém o padrão visual definido, se os dados exibidos são consistentes com a OS e se a validação pública por QR Code reduz o risco de falsificação.",
    )
    add_paragraph(
        document,
        "Registre as observações neste próprio documento, marcando cada item como OK, Parcial ou Erro. Evite concentrar ocorrências no grupo de WhatsApp; o documento preenchido facilita a triagem e a priorização técnica.",
    )

    add_heading(document, "1. Dados da rodada")
    add_info_table(document)

    add_heading(document, "2. Fluxo esperado")
    for item in [
        "Entrar no ambiente de homologação com o perfil de Qualidade ou responsável técnico.",
        "Acessar Certificados e Histórico.",
        "Abrir um certificado emitido ou uma OS encerrada com certificado disponível.",
        "Gerar ou visualizar o certificado em PDF.",
        "Conferir layout, dados dinâmicos, QR Code, código curto, SHA-256, fotos e assinatura.",
        "Abrir a validação pública pelo QR Code e comparar os dados com o PDF.",
        "Registrar no roteiro qualquer divergência visual, de dados, de fluxo ou de leitura do QR Code.",
    ]:
        add_bullet(document, item)

    add_heading(document, "3. Checklist principal")
    add_result_table(document)

    add_heading(document, "4. Teste dedicado do QR Code")
    add_paragraph(
        document,
        "Este bloco é obrigatório. O QR Code deve funcionar tanto na tela quanto no impresso. Sempre compare se a rota pública abriu exatamente o certificado correspondente, sem trocar cliente, OS, serviço, data ou hash.",
    )
    add_qr_validation_table(document)

    add_sample_image(document)

    add_heading(document, "6. Registro das observações")
    add_response_area(document)

    add_heading(document, "7. Mensagem para o grupo")
    add_paragraph(
        document,
        "Pessoal, segue o roteiro para validar os certificados do Atenza FieldOps em homologação. Por favor, registrem qualquer OK, erro, divergência visual, problema de QR Code ou sugestão diretamente no DOCX. Evitem mandar ocorrências soltas no grupo, para não perdermos o histórico e conseguirmos priorizar corretamente os ajustes.",
        indent=False,
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
