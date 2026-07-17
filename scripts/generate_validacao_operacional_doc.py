from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "docs"
OUTPUT_FILE = OUTPUT_DIR / "roteiro_validacao_operacional_ciperprag_hub.docx"
LOGO_FILE = ROOT / "src" / "assets" / "logo_ciperprag.png"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
      tc_w = OxmlElement("w:tcW")
      tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    tc_w.set(qn("w:type"), "dxa")


def configure_run(run, *, size=11, bold=False, color="000000", font="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)


def style_paragraph(paragraph, *, size=11, bold=False, color="000000", font="Calibri", space_after=6, space_before=0, line=1.15):
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        configure_run(run, size=size, bold=bold, color=color, font=font)


def add_title(document: Document) -> None:
    if LOGO_FILE.exists():
        logo = document.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_run = logo.add_run()
        logo_run.add_picture(str(LOGO_FILE), width=Inches(2.25))
        logo.paragraph_format.space_after = Pt(8)

    title = document.add_paragraph()
    title_run = title.add_run("Roteiro de Validacao Operacional")
    configure_run(title_run, size=23, bold=False, color="111111", font="Arial")
    title.paragraph_format.space_after = Pt(2)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("Ciperprag Hub | Homologacao da area Operacional")
    configure_run(subtitle_run, size=12, bold=False, color="4B5563", font="Arial")
    subtitle.paragraph_format.space_after = Pt(10)

    chip = document.add_paragraph()
    chip_run = chip.add_run("Ambiente produtivo: http://89.116.214.65:3010 | Sem login | Base PostgreSQL em uso")
    configure_run(chip_run, size=10.5, bold=False, color="1F2937", font="Arial")
    chip.paragraph_format.space_after = Pt(12)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[f"Heading {level}"]
    run = paragraph.add_run(text)
    if level == 1:
        configure_run(run, size=16, bold=False, color="111111", font="Arial")
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(6)
    elif level == 2:
        configure_run(run, size=13, bold=False, color="111111", font="Arial")
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(4)
    else:
        configure_run(run, size=11.5, bold=True, color="374151", font="Arial")
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(3)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    style_paragraph(paragraph, size=10.8, font="Arial", space_after=6)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        style_paragraph(paragraph, size=10.8, font="Arial", space_after=3)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)
        style_paragraph(paragraph, size=10.8, font="Arial", space_after=3)


def add_info_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for label, value in rows:
        row = table.add_row().cells
        set_cell_width(row[0], 4.2)
        set_cell_width(row[1], 11.8)
        row[0].text = label
        row[1].text = value
        for index, cell in enumerate(row):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, size=10.5, bold=index == 0, color="111111", font="Arial", space_after=0)
            if index == 0:
                set_cell_shading(cell, "F3F4F6")
    document.add_paragraph()


def add_checklist_table(document: Document, rows: list[tuple[str, str, str, str]]) -> None:
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    headers = ["ID", "Cenario", "Como validar", "Resultado esperado", "Status"]
    widths = [1.1, 3.2, 6.0, 6.0, 1.7]
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = headers[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "DDEFE2")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            style_paragraph(paragraph, size=9.8, bold=True, color="17311F", font="Arial", space_after=0)

    for row_data in rows:
        cells = table.add_row().cells
        values = [row_data[0], row_data[1], row_data[2], row_data[3], ""]
        for idx, value in enumerate(values):
            cells[idx].text = value
            set_cell_width(cells[idx], widths[idx])
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cells[idx].paragraphs:
                style_paragraph(paragraph, size=9.7, color="111111", font="Arial", space_after=0, line=1.12)
    document.add_paragraph()


def add_issue_log_table(document: Document) -> None:
    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    headers = ["Item", "Modulo", "Descricao da divergencia", "Evidencia", "Acao"]
    widths = [1.2, 3.0, 7.0, 3.5, 3.3]
    for idx, cell in enumerate(table.rows[0].cells):
        cell.text = headers[idx]
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "F3F4F6")
        for paragraph in cell.paragraphs:
            style_paragraph(paragraph, size=9.8, bold=True, color="111111", font="Arial", space_after=0)
    for item in range(1, 5):
        row = table.add_row().cells
        values = [str(item), "", "", "", ""]
        for idx, value in enumerate(values):
            row[idx].text = value
            set_cell_width(row[idx], widths[idx])
            for paragraph in row[idx].paragraphs:
                style_paragraph(paragraph, size=9.7, color="111111", font="Arial", space_after=10)
    document.add_paragraph()


def add_footer(document: Document) -> None:
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.text = "Ciperprag Hub | Roteiro de validacao operacional | Junho/2026"
    style_paragraph(paragraph, size=8.5, color="6B7280", font="Arial", space_after=0)


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 3"].font.name = "Arial"

    add_title(document)

    add_heading(document, "1. Objetivo e escopo", 1)
    add_body(document, "Este roteiro foi preparado para a equipe operacional validar o fluxo principal do Ciperprag Hub em producao, considerando somente os modulos operacionais nesta etapa.")
    add_bullets(
        document,
        [
            "Escopo atual: Dashboard, Agendamentos, Ordens de Servico, encerramento de OS, Certificados e Historico, Medicao e sugestao de recorrencia.",
            "Ambiente de validacao: VPS publica em http://89.116.214.65:3010.",
            "Autenticacao: nao se aplica nesta versao, pois o sistema opera sem login.",
            "Base utilizada: dados do banco PostgreSQL em producao, sem uso de mocks locais.",
        ],
    )

    add_heading(document, "2. Informacoes rapidas para acesso", 1)
    add_info_table(
        document,
        [
            ("URL", "http://89.116.214.65:3010"),
            ("Login", "Nao possui login nesta fase."),
            ("Perfil alvo", "Equipe operacional e administracao interna."),
            ("Area a validar", "Menu Operacional: Dashboard, Agendamentos, Ordens de Servico, Certificados e Historico, Medicao."),
            ("Observacao", "Sempre validar com internet estavel e anexos reais quando houver fechamento de OS."),
        ],
    )

    add_heading(document, "3. Dados uteis ja disponiveis no ambiente", 1)
    add_info_table(
        document,
        [
            ("Agendamento de referencia", "AG-MQPCLRC5"),
            ("OS de referencia", "OS-2677"),
            ("Certificado de referencia", "HSH-2026-9FY7"),
            ("Recorrencia sugerida", "Novo agendamento sugerido para 22/07/2026 apos conclusao do servico."),
            ("Observacao", "Se esses registros forem alterados ao longo do uso, a equipe pode criar novos registros e seguir a mesma logica de validacao."),
        ],
    )

    add_heading(document, "4. Sequencia recomendada de homologacao", 1)
    add_numbered(
        document,
        [
            "Acessar Dashboard e confirmar que os indicadores estao carregando com dados reais.",
            "Criar ou editar um agendamento com equipe designada, tecnicos e veiculo.",
            "Gerar a Ordem de Servico a partir do agendamento.",
            "Imprimir a OS e conferir se o layout e os campos dinamicos estao preenchidos.",
            "Encerrar a OS com informacoes de campo, tag do equipamento e ate tres fotos.",
            "Se o servico permitir, emitir o certificado e conferir o documento final.",
            "Verificar no Historico se a execucao aparece para o cliente.",
            "Ir para Medicao, filtrar o periodo, gerar a medicao e emitir o PDF.",
            "Confirmar a sugestao de recorrencia quando o servico for recorrente.",
        ],
    )

    add_heading(document, "5. Checklist por modulo", 1)
    add_heading(document, "5.1 Dashboard", 2)
    add_checklist_table(
        document,
        [
            ("D-01", "Indicadores principais", "Abrir a tela inicial e observar cards, agenda e visoes de apoio.", "Cards e listagens devem abrir com dados reais, sem erro visual ou travamento."),
            ("D-02", "Atalhos operacionais", "Usar os links rapidos para navegar para Agendamentos, OS, Certificados e Medicao.", "A navegacao deve ser direta, sem recarregamento quebrado e sem menu ilegivel."),
        ],
    )

    add_heading(document, "5.2 Agendamentos", 2)
    add_checklist_table(
        document,
        [
            ("A-01", "Criar agendamento", "Selecionar cliente, contrato, servico, data e local de execucao.", "O agendamento deve ser salvo com status correto e aparecer nas listagens."),
            ("A-02", "Designar equipe", "No mesmo agendamento, informar tecnicos, veiculo e dados complementares.", "Tecnicos e veiculo devem ficar vinculados ao agendamento e seguir para a OS."),
            ("A-03", "Editar agendamento", "Abrir um agendamento ja salvo, alterar campos e salvar novamente.", "As alteracoes devem persistir no banco e refletir nas proximas telas."),
        ],
    )

    add_heading(document, "5.3 Ordens de Servico", 2)
    add_checklist_table(
        document,
        [
            ("OS-01", "Gerar OS", "A partir do agendamento, acionar a geracao da Ordem de Servico.", "A OS deve receber numero, manter vinculo com o agendamento e ir para a tela de gestao."),
            ("OS-02", "Gerenciar OS existentes", "Acessar a lista de OS, abrir, editar e localizar uma ordem ja criada.", "A listagem deve exibir todas as OS cadastradas e permitir continuidade do fluxo."),
            ("OS-03", "Impressao da OS", "Usar a opcao de imprimir na OS gerada.", "O documento deve abrir com layout proximo ao modelo de referencia, incluindo campos dinamicos, equipe, cliente, observacoes e blocos de assinatura."),
        ],
    )

    add_heading(document, "5.4 Encerramento de OS", 2)
    add_checklist_table(
        document,
        [
            ("F-01", "Finalizar ordem", "Marcar a OS como concluida ao retorno da equipe de campo.", "A OS deve mudar de status e permanecer consultavel."),
            ("F-02", "Registrar tag", "Informar a tag do equipamento quando o servico exigir rastreabilidade.", "A tag deve ficar salva e disponivel para certificado e historico."),
            ("F-03", "Anexar fotos", "Enviar ate tres fotos de evidencia na finalizacao.", "As imagens devem ser aceitas sem travar a tela e vinculadas ao fechamento."),
            ("F-04", "Observacoes finais", "Preencher conclusoes, pendencias ou observacoes da equipe.", "As informacoes devem ficar registradas para consultas futuras."),
        ],
    )

    add_heading(document, "5.5 Certificados e Historico", 2)
    add_checklist_table(
        document,
        [
            ("C-01", "Emitir certificado", "Concluir uma OS elegivel e gerar o certificado do servico.", "O certificado deve respeitar o template dinamico, preencher dados do cliente, servico, data e responsavel."),
            ("C-02", "Historico do cliente", "Abrir a aba Historico e localizar servicos do cliente, com ou sem certificado.", "A linha do tempo deve exibir execucoes concluidas independentemente da emissao do certificado."),
        ],
    )

    add_heading(document, "5.6 Medicao", 2)
    add_checklist_table(
        document,
        [
            ("M-01", "Baixa contratual", "Concluir um servico e acessar Medicao no periodo correspondente.", "O item executado deve compor a medicao e refletir baixa no contrato."),
            ("M-02", "Filtro por intervalo", "Definir um intervalo de datas e gerar a medicao.", "A tela deve listar somente as OS encerradas dentro do periodo informado."),
            ("M-03", "PDF da medicao", "Usar a opcao de imprimir ou gerar PDF.", "O PDF deve ser emitido com cabecalho, itens, totais e periodo de medicao."),
        ],
    )

    add_heading(document, "5.7 Recorrencia", 2)
    add_checklist_table(
        document,
        [
            ("R-01", "Sugestao automatica", "Concluir um servico recorrente e voltar para Agendamentos.", "O sistema deve sugerir uma nova data com base na recorrencia do servico ou contrato."),
            ("R-02", "Confirmacao da sugestao", "Confirmar a data sugerida na interface.", "Ao confirmar, um novo agendamento deve ser criado e voltar para o inicio do fluxo operacional."),
        ],
    )

    add_heading(document, "6. Criterios minimos de aceite", 1)
    add_bullets(
        document,
        [
            "Nenhuma etapa principal pode depender de dados mockados locais.",
            "Todas as gravacoes devem persistir no banco e reaparecer apos atualizar a pagina.",
            "Menu lateral deve estar legivel e navegavel em tela desktop e mobile.",
            "OS, certificado e medicao devem abrir para impressao sem quebra de layout critica.",
            "Fluxos com recorrencia devem sugerir novo agendamento sem duplicidade indevida.",
        ],
    )

    add_heading(document, "7. Registro de divergencias", 1)
    add_body(document, "Use a grade abaixo para anotar qualquer comportamento inesperado durante a homologacao.")
    add_issue_log_table(document)

    add_heading(document, "8. Parecer final", 1)
    add_body(document, "Ao fim da validacao, registrar se a area Operacional esta aprovada para uso assistido, aprovada com ressalvas ou reprovada para correcao adicional.")
    add_info_table(
        document,
        [
            ("Resultado da rodada", ""),
            ("Responsavel pela validacao", ""),
            ("Data", ""),
            ("Observacoes finais", ""),
        ],
    )

    add_footer(document)
    return document


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    main()
