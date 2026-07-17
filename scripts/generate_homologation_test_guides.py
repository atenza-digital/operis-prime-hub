# -*- coding: utf-8 -*-
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(os.environ["ATENZA_TEMPLATE_DOCX"])
OUT = ROOT / "docs" / "cliente" / "homologacao_roteiros"
PRINTS = ROOT / "docs" / "evidencias" / "etapa7_homologacao" / "prints_roteiros"

VERSION = "0.6.3"
URL = "http://89.116.214.65:3010"
DATA = "07/07/2026"
PRAZO = "Pendente de confirmação"
RESP_ATENZA = "Ewerton Gomes Almeida"
CLIENTE = "Ciperprag"
PROJETO = "Atenza FieldOps"

CREDS = {
    "comercial": (os.environ["HML_COMERCIAL_USER"], os.environ["HML_COMERCIAL_PASS"]),
    "operacao": (os.environ["HML_OPERACAO_USER"], os.environ["HML_OPERACAO_PASS"]),
    "qualidade": (os.environ["HML_QUALIDADE_USER"], os.environ["HML_QUALIDADE_PASS"]),
    "medicao": (os.environ["HML_MEDICAO_USER"], os.environ["HML_MEDICAO_PASS"]),
}

PROFILES = {
    "comercial": {
        "title": "Roteiro de Testes - Perfil Comercial",
        "subtitle": "Validação de clientes, serviços, proposta comercial, aprovação e geração de contrato.",
        "status": "Em teste",
        "criteria": "Proposta criada, visualizada, aprovada e convertida em contrato vigente integrado ao operacional.",
        "screens": [("Tela de login", "01-login.png"), ("Contratos, propostas e contrato do cliente", "03-comercial-contratos.png")],
        "steps": [
            ("CT-COM-001", "Acessar ambiente de homologação", "Entrar na URL de homologação com o usuário comercial informado neste roteiro.", "Sistema deve abrir o painel permitido para o perfil, sem solicitar decisões técnicas."),
            ("CT-COM-002", "Conferir cliente e serviços", "Acessar Comercial, conferir se existe cliente de teste e se os serviços/produtos estão cadastrados.", "Cliente e serviços devem estar disponíveis para criação de proposta."),
            ("CT-COM-003", "Criar proposta", "Criar uma nova proposta com cliente, itens, quantidades, valores, vigência, periodicidade e condição comercial.", "Proposta deve salvar sem erro e exibir dados principais de forma clara."),
            ("CT-COM-004", "Visualizar/imprimir proposta", "Abrir a visualização/impressão da proposta gerada.", "Documento deve apresentar logo, cliente, itens, valores, condições e texto sem cortes."),
            ("CT-COM-005", "Aprovar proposta", "Alterar o status da proposta para aprovado.", "Sistema deve deixar claro que a proposta está aprovada e pronta para gerar contrato."),
            ("CT-COM-006", "Gerar contrato", "Gerar contrato a partir da proposta aprovada.", "Contrato deve nascer vigente ou pronto para vigência e liberar itens operacionais para agendamento."),
            ("CT-COM-007", "Contrato do cliente", "Testar o caminho alternativo quando o cliente já possui modelo próprio.", "Contrato do cliente deve liberar saldo operacional sem exigir proposta duplicada."),
        ],
        "data": ["Cliente de teste existente na base", "Serviços/produtos cadastrados no catálogo", "Valores e condições comerciais fictícios para homologação"],
    },
    "operacao": {
        "title": "Roteiro de Testes - Perfil Operacional",
        "subtitle": "Validação de agendamento, equipe, veículo, geração de OS, impressão e encerramento com evidências.",
        "status": "Em teste",
        "criteria": "Agendamento criado a partir de contrato com saldo, OS gerada, impressa e encerrada com evidências.",
        "screens": [("Tela de login", "01-login.png"), ("Agendamentos", "04-agendamentos.png"), ("Ordens de Serviço", "05-ordens-servico.png")],
        "steps": [
            ("CT-OP-001", "Acessar agendamentos", "Entrar na URL de homologação com o usuário operacional e abrir Agendamentos.", "Tela deve orientar o fluxo sem exigir conhecimento técnico."),
            ("CT-OP-002", "Selecionar contrato com saldo", "Escolher cliente, contrato/serviço e confirmar o saldo operacional exibido.", "Apenas contratos vigentes com saldo devem aparecer."),
            ("CT-OP-003", "Informar execução", "Preencher data, local de execução, técnicos, veículo e tag/equipamento quando aplicável.", "Dados devem ficar salvos e serem levados para a OS."),
            ("CT-OP-004", "Gerar OS", "Gerar ordem de serviço a partir do agendamento criado.", "OS deve receber numeração automática e dados do agendamento."),
            ("CT-OP-005", "Imprimir via da equipe", "Abrir a impressão da OS.", "OS deve estar legível, com campos dinâmicos e sem texto cortado."),
            ("CT-OP-006", "Encerrar OS", "Registrar retorno de campo: data de execução, quantidade executada, tag/equipamento, observações e até 3 fotos.", "OS deve mudar para encerrada e entrar no histórico, certificado quando aplicável e medição."),
            ("CT-OP-007", "Recorrência", "Quando o serviço for recorrente, confirmar a sugestão de próximo agendamento.", "Sistema deve criar novo agendamento e reiniciar o fluxo operacional."),
        ],
        "data": ["Contrato vigente com saldo", "Equipe/técnicos disponíveis", "Veículo quando aplicável", "Tag/equipamento quando o serviço exigir", "Até 3 fotos de teste"],
    },
    "qualidade": {
        "title": "Roteiro de Testes - Perfil Qualidade / Responsável Técnico",
        "subtitle": "Validação de certificado, fotos da OS, QR Code, consulta pública e histórico de serviços.",
        "status": "Em teste",
        "criteria": "Certificado gerado com dados corretos, QR Code funcional e histórico mostrando serviços com e sem certificado.",
        "screens": [("Tela de login", "01-login.png"), ("Certificados e Histórico", "06-certificados-historico.png"), ("Auditoria de anexos", "08-auditoria-anexos.png")],
        "steps": [
            ("CT-QL-001", "Acessar certificados", "Entrar com o usuário de qualidade e abrir Certificados e Histórico.", "Tela deve listar certificados e histórico conforme permissões."),
            ("CT-QL-002", "Localizar OS encerrada", "Encontrar uma OS encerrada cujo serviço permita certificado.", "OS elegível deve estar disponível para emissão de certificado."),
            ("CT-QL-003", "Gerar certificado", "Gerar certificado do serviço executado.", "Certificado deve trazer cliente, CNPJ, serviço, OS, data, técnico, local, tag e fotos quando existirem."),
            ("CT-QL-004", "Validar fotos", "Confirmar se até 3 fotos anexadas na OS aparecem no certificado quando aplicável.", "Fotos devem ser dinâmicas e vinculadas à OS, não fixas no documento."),
            ("CT-QL-005", "Validar QR Code", "Ler o QR Code ou abrir a rota pública de validação.", "Consulta pública deve retornar dados compatíveis com o certificado."),
            ("CT-QL-006", "Ver histórico", "Abrir histórico do cliente/serviços.", "Histórico deve mostrar serviços que geraram certificado e serviços que não geraram."),
            ("CT-QL-007", "Auditar anexos", "Conferir documento/anexo histórico quando disponível.", "Anexo deve exibir hash, categoria, imutabilidade e opção de download quando permitido."),
        ],
        "data": ["OS encerrada com serviço que permite certificado", "Fotos anexadas na OS", "Tag/equipamento quando aplicável", "QR Code do certificado gerado"],
    },
    "medicao": {
        "title": "Roteiro de Testes - Perfil Medição / Administrativo",
        "subtitle": "Validação de medição operacional, consolidação de OS, NF enviada, cobrança, pagamento e baixa manual no ERP.",
        "status": "Em teste",
        "criteria": "Medição gerada por período, com OS encerradas, total correto e acompanhamento financeiro operacional atualizado.",
        "screens": [("Tela de login", "01-login.png"), ("Medição", "07-medicao.png")],
        "steps": [
            ("CT-MED-001", "Acessar medição", "Entrar com o usuário de medição e abrir a tela Medição.", "Tela deve deixar claro que acompanha medição, NF e pagamento, sem substituir o ERP."),
            ("CT-MED-002", "Selecionar cliente e período", "Escolher cliente e intervalo de datas com OS encerradas.", "Devem aparecer apenas OS encerradas, ainda não medidas e dentro do período."),
            ("CT-MED-003", "Gerar medição", "Gerar a medição do período selecionado.", "Medição deve consolidar itens, quantidades, valores e total geral."),
            ("CT-MED-004", "Visualizar PDF", "Abrir a impressão/PDF da medição.", "Documento deve estar em A4 retrato, institucional, com logo, dados da contratada, cliente, itens e assinaturas."),
            ("CT-MED-005", "Atualizar NF", "Informar número da NF, data de envio e status NF enviada.", "Kanban/status deve refletir que a NF foi enviada."),
            ("CT-MED-006", "Acompanhar cobrança", "Alterar para aguardando pagamento quando houver cobrança pendente.", "Sistema deve facilitar acompanhamento sem virar contas a receber completo."),
            ("CT-MED-007", "Baixa manual no ERP", "Atualizar para pago no ERP após baixa manual realizada fora do FieldOps.", "Medição deve indicar que foi paga/baixada no ERP."),
        ],
        "data": ["Cliente com OS encerradas no período", "Período de teste informado pela equipe", "Número de NF fictício para homologação", "Status financeiro operacional"],
    },
}

WHATSAPP = """Pessoal, liberamos a rodada de testes da homologação do Atenza FieldOps para validar principalmente o fluxo de OS e documentos.

Cada perfil recebeu um roteiro próprio em DOCX com URL, usuário, senha de teste, passos e resultado esperado.

Importante: por favor, não reportem erros ou observações soltas aqui no grupo. Registrem tudo no próprio documento do roteiro, com print, número da OS/proposta/certificado/medição e uma observação curta do que aconteceu. Assim conseguimos mapear, priorizar e corrigir sem perder informação.

Ao finalizar, enviem os documentos preenchidos de volta para consolidação da homologação."""


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def text_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(str(text))
    run.bold = bold


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def para(doc: Document, text: str = "", style: str | None = None, justify: bool = True, indent: bool = True, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0.7) if indent else Cm(0)
    run = paragraph.add_run(text)
    run.bold = bold
    return paragraph


def heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.add_run(text)
    return paragraph


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = tbl.rows[0].cells[index]
        text_cell(cell, header, True)
        shade(cell, "F2F2ED")
    for row in rows:
        cells = tbl.add_row().cells
        for index, value in enumerate(row):
            text_cell(cells[index], value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in tbl.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Cm(width)
    doc.add_paragraph()
    return tbl


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.add_run(item)


def image(doc: Document, caption: str, filename: str) -> None:
    path = PRINTS / filename
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.add_run().add_picture(str(path), width=Inches(6.2))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.first_line_indent = Cm(0)
    caption_paragraph.add_run(caption).italic = True


def build_doc(profile_key: str, info: dict) -> Path:
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    para(doc, "DOCUMENTO INSTITUCIONAL", justify=False, indent=False, bold=True)
    heading(doc, info["title"], 1)
    para(doc, info["subtitle"])
    heading(doc, "Ficha do projeto", 2)
    table(doc, ["Campo", "Valor"], [
        ["Projeto", PROJETO],
        ["Cliente final", CLIENTE],
        ["Responsável Atenza", RESP_ATENZA],
        ["Responsável pela homologação", "Equipe Ciperprag"],
        ["Ambiente testado", "Homologação"],
        ["URL", URL],
        ["Versão/build", VERSION],
        ["Data da rodada", DATA],
        ["Prazo de retorno", PRAZO],
        ["Status da validação", info["status"]],
        ["Critério de aprovação", info["criteria"]],
    ], [5, 11])
    heading(doc, "1. Acesso de teste", 2)
    email, senha = CREDS[profile_key]
    table(doc, ["Campo", "Informação"], [
        ["URL de homologação", URL],
        ["Usuário", email],
        ["Senha de teste", senha],
        ["Perfil de teste", profile_key.capitalize()],
        ["Observação", "Conta exclusiva para homologação. Não usar em produção."],
    ], [5, 11])
    para(doc, "Use somente o ambiente de homologação indicado neste roteiro. Não registre dados reais sensíveis; use dados fictícios quando precisar preencher observações, NF ou anexos.")
    heading(doc, "2. Como registrar o resultado", 2)
    para(doc, "Durante o teste, preencha a tabela de resultado ao final deste roteiro. Quando encontrar problema, registre o passo, descreva o ocorrido em uma frase objetiva, informe o número do documento gerado e cole um print no próprio arquivo antes de devolver.")
    table(doc, ["Status", "Quando usar"], [
        ["Aprovado", "O passo funcionou e o resultado esperado apareceu sem dúvida para o usuário."],
        ["Aprovado com ressalva", "Funcionou, mas houve confusão, texto ruim, acento errado, excesso de cliques ou detalhe visual."],
        ["Reprovado", "Impediu seguir o fluxo, gerou dado incorreto, documento errado ou erro na tela."],
    ], [4, 12])
    table(doc, ["Severidade", "Critério"], [
        ["Alta", "Bloqueia o fluxo ou gera documento/dado incorreto."],
        ["Média", "Permite contorno, mas confunde o usuário ou exige retrabalho."],
        ["Baixa", "Ajuste visual, texto, acentuação, alinhamento ou melhoria de conveniência."],
    ], [4, 12])
    heading(doc, "3. Dados necessários para esta rodada", 2)
    bullets(doc, info["data"])
    heading(doc, "4. Roteiro de teste", 2)
    table(doc, ["ID", "Cenário", "Passos", "Resultado esperado", "Status"], [[*step, "Pendente"] for step in info["steps"]], [2.2, 3.4, 4.8, 4.8, 2])
    heading(doc, "5. Prints de referência", 2)
    para(doc, "Os prints abaixo servem apenas para ajudar a localizar a tela. A aparência pode variar um pouco conforme tamanho do monitor ou navegador, mas o fluxo e os nomes principais devem permanecer equivalentes.")
    for caption, filename in info["screens"]:
        image(doc, caption, filename)
    heading(doc, "6. Registro de divergências", 2)
    para(doc, "Preencha uma linha para cada problema encontrado. Se o teste for aprovado sem ressalvas, registre 'Sem divergências' na primeira linha.")
    table(doc, ["ID", "Passo", "O que aconteceu", "Documento/Número", "Print anexado?", "Severidade", "Status"], [
        ["HML-001", "", "", "", "Sim/Não", "", "Aberto"],
        ["HML-002", "", "", "", "Sim/Não", "", "Aberto"],
        ["HML-003", "", "", "", "Sim/Não", "", "Aberto"],
    ], [2, 2.2, 5, 3, 2.2, 2, 2])
    heading(doc, "7. Resultado final do perfil", 2)
    table(doc, ["Campo", "Preenchimento pela equipe"], [
        ["Resultado geral", "Aprovado / Aprovado com ressalva / Reprovado"],
        ["Nome de quem testou", ""],
        ["Data do teste", ""],
        ["Resumo das principais observações", ""],
        ["Pode seguir para próxima etapa?", "Sim / Não"],
    ], [5, 11])
    heading(doc, "Histórico de versões", 2)
    table(doc, ["Versão", "Data", "Responsável", "Descrição"], [["v1.0", DATA, RESP_ATENZA, f"Emissão inicial do roteiro de testes para o perfil {profile_key}."]], [2, 3, 5, 7])
    out = OUT / f"Roteiro_Testes_Atenza_FieldOps_{profile_key.capitalize()}_v1.0.docx"
    doc.save(out)
    return out


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    outputs = [build_doc(key, info) for key, info in PROFILES.items()]
    (OUT / "MENSAGEM_WHATSAPP_HOMOLOGACAO.md").write_text(WHATSAPP, encoding="utf-8")
    for output in outputs:
        print(output)
    print(OUT / "MENSAGEM_WHATSAPP_HOMOLOGACAO.md")


if __name__ == "__main__":
    main()
