# -*- coding: utf-8 -*-
"""Build the focused regression guide for the homologation retest."""
from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(
    r"C:\Projetos\Documentações\02_Templates\01_Operacionais\Template_Roteiro_Testes_Validacao_Atenza.docx"
)
REPO_OUT = ROOT / "docs" / "cliente" / "homologacao_roteiros" / "Roteiro_Validacao_Regressao_Atenza_FieldOps_Tarcisio_v1.5.docx"
DOWNLOAD_OUT = Path(r"C:\Users\herto\Downloads\Roteiro_Validacao_Regressao_Atenza_FieldOps_Tarcisio_v1.5.docx")
WHATSAPP_OUT = ROOT / "docs" / "cliente" / "homologacao_roteiros" / "MENSAGEM_WHATSAPP_HOMOLOGACAO_TARCISIO_v1.5.md"

URL = "https://fieldops-homologacao.atenza.digital/login"
VERSION = "0.6.3"
DEPLOY_COMMIT = "2c887c3"
WORKFLOW = "29842924268"
ACCENT = "087F5B"
INK = "142235"
MUTED = "5B6878"
LIGHT = "F4F7F6"


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, *, size: float = 9.5, bold: bool = False, color: str = INK) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc: Document, text: str = "", *, size: float = 9.5, bold: bool = False, color: str = INK, indent: bool = False, after: float = 6, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0.7) if indent else Cm(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(13 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run(r, size=18 if level == 1 else 13 if level == 2 else 11, bold=True, color=ACCENT if level > 1 else INK)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.25)
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(text), size=9.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, ACCENT)
        set_cell_margins(cell, 110, 120, 110, 120)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(header), size=8.3, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            set_cell_shading(cell, "FFFFFF" if len(table.rows) % 2 else LIGHT)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run(p.add_run(str(value)), size=8.1)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    add_paragraph(doc, after=2)
    return table


def add_image(doc: Document, relative: str, caption: str) -> None:
    path = ROOT / relative
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(5.8))
    cap = add_paragraph(doc, caption, size=8, color=MUTED, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.runs[0].italic = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Atenza FieldOps • Homologação • Roteiro de revalidação v1.5"
    for run in footer.runs:
        set_run(run, size=7.5, color=MUTED)


def build() -> None:
    REPO_OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    configure_document(doc)

    add_paragraph(doc, "DOCUMENTO DE HOMOLOGAÇÃO", size=9, bold=True, color=ACCENT, after=4)
    add_heading(doc, "Roteiro de Revalidação do Atenza FieldOps", 1)
    add_paragraph(doc, "Rodada dirigida após os ajustes apontados no roteiro de testes do Tarcísio", size=11, color=MUTED, after=10)
    add_paragraph(doc, "Objetivo: confirmar somente as correções publicadas na homologação, sem repetir toda a matriz já aprovada. Registre cada resultado neste arquivo e devolva o DOCX preenchido à Atenza.", indent=True)

    add_heading(doc, "1. Acesso e identificação", 2)
    add_table(doc, ["Campo", "Informação"], [
        ["URL", URL],
        ["Ambiente", "Homologação — uso exclusivo para testes"],
        ["Versão exibida", VERSION],
        ["Deploy validado", f"Commit {DEPLOY_COMMIT} • workflow CI/CD {WORKFLOW}"],
        ["Perfil recomendado", "Usar o mesmo perfil do caso testado no roteiro v1.4"],
        ["Senha", "Usar a senha temporária enviada pelo canal seguro. Não registrar no documento."],
    ], [4.0, 12.0])
    add_paragraph(doc, "Antes de iniciar, confirme que o endereço é exatamente o domínio HTTPS acima. Não use o IP antigo nem a porta local.", bold=True, color=ACCENT, indent=True)

    add_heading(doc, "2. Como registrar o resultado", 2)
    add_bullet(doc, "Aprovado: o comportamento corresponde ao resultado esperado.")
    add_bullet(doc, "Aprovado com ressalva: funciona, mas há uma diferença visual ou de usabilidade que deve ser registrada.")
    add_bullet(doc, "Reprovado: há erro, bloqueio, dado divergente, documento incorreto ou QR Code que não valida.")
    add_bullet(doc, "Para cada problema, informe o ID, a tela/documento, número da proposta/contrato/OS/certificado e cole um print no próprio DOCX.")
    add_bullet(doc, "Não envie ocorrências soltas no grupo de WhatsApp. O grupo serve apenas para avisar que a rodada foi concluída.")

    add_heading(doc, "3. Regressão dirigida", 2)
    rows = [
        ["RET-01", "Dashboard", "Abrir o dashboard em uma janela de desktop e reduzir a largura até o breakpoint de tablet.", "Os cinco cards de Fluxo recomendado mantêm largura, espaçamento, textos e marcadores sem sobreposição ou deformação.", "____"],
        ["RET-02", "Comercial", "Abrir Contratos e conferir a ordenação depois de criar ou localizar dois contratos.", "O contrato criado mais recentemente aparece no topo; a ordem permanece estável após atualizar a tela.", "____"],
        ["RET-03", "Proposta", "Abrir o PDF da proposta e ir à página de aceite.", "As duas assinaturas ficam na mesma linha, alinhadas horizontalmente, sem corte, sobreposição ou empresa em terceira linha.", "____"],
        ["RET-04", "Agenda", "Abrir um agendamento existente e visualizar os detalhes.", "Os detalhes exibem claramente a equipe designada e o veículo designado, com dados coerentes com o agendamento.", "____"],
        ["RET-05", "OS e certificados", "Encerrar uma OS com mais de uma TAG/equipamento aplicável, usando os dados permitidos no formulário.", "É criado um certificado para cada TAG, cada um com sua própria identificação, hash, QR Code e vínculo correto com a mesma OS.", "____"],
        ["RET-06", "Certificado", "Abrir cada certificado gerado e comparar cliente, serviço, TAG, data, fotos e responsável com a OS.", "Não há mistura de cliente, serviço, TAG, texto declaratório, fotos ou responsável entre certificados.", "____"],
        ["RET-07", "QR Code", "Ler o QR Code de cada certificado na tela e em uma impressão A4.", "O QR abre a URL pública HTTPS do certificado correspondente e apresenta status, tenant, número, OS, serviço, execução, validade e SHA-256.", "____"],
    ]
    add_table(doc, ["ID", "Área", "Ação", "Resultado esperado", "Status"], rows, [1.35, 2.0, 5.2, 6.1, 1.35])

    add_heading(doc, "4. Evidência visual de referência", 2)
    add_paragraph(doc, "Os prints abaixo são apenas orientação de navegação. A aprovação deve ser feita na aplicação publicada, com os dados atuais de homologação.")
    add_image(doc, "docs/evidencias/etapa7_homologacao/prints_roteiros/02-dashboard.png", "Dashboard: conferir principalmente os cards do fluxo recomendado.")
    add_image(doc, "docs/evidencias/etapa7_homologacao/prints_roteiros/03-comercial-contratos.png", "Comercial: conferir a lista e a abertura dos documentos.")
    add_image(doc, "docs/evidencias/etapa7_homologacao/prints_roteiros/04-agendamentos.png", "Agenda: abrir o detalhe de um item e conferir equipe e veículo.")
    add_image(doc, "docs/evidencias/etapa7_homologacao/prints_roteiros/05-ordens-servico.png", "OS: conferir encerramento e evidências antes de gerar certificados.")
    add_image(doc, "docs/evidencias/etapa7_homologacao/prints_roteiros/06-certificados-historico.png", "Certificados e histórico: localizar certificados por OS e TAG.")

    add_heading(doc, "5. Teste real do QR Code", 2)
    add_paragraph(doc, "Este teste é obrigatório para encerrar a pendência QL-03/FINAL-05. Use dois celulares diferentes. Primeiro leia o QR diretamente no PDF em tela; depois imprima o mesmo PDF em A4 e leia novamente.", indent=True)
    add_table(doc, ["Verificação", "Preenchimento"], [
        ["Certificado 1 / código", "____________________________________________"],
        ["Certificado 2 / código, se houver", "____________________________________________"],
        ["Celular/aplicativo usado 1", "____________________________________________"],
        ["Celular/aplicativo usado 2", "____________________________________________"],
        ["Leitura em tela", "Aprovado / Ressalva / Reprovado"],
        ["Leitura no papel A4", "Aprovado / Ressalva / Reprovado"],
        ["URL aberta", "____________________________________________"],
        ["Observação", "____________________________________________\n____________________________________________"],
    ], [5.3, 11.7])
    add_paragraph(doc, "Critério: não basta o QR abrir uma página. A página deve corresponder ao certificado lido, sem trocar tenant, cliente, serviço, OS, TAG ou hash.", bold=True, color=ACCENT, indent=True)

    add_heading(doc, "6. Registro de divergências", 2)
    add_table(doc, ["ID", "Tela/documento", "O que aconteceu", "Evidência", "Severidade", "Status"], [
        ["HML-RET-01", "", "", "", "", "Aberto"],
        ["HML-RET-02", "", "", "", "", "Aberto"],
        ["HML-RET-03", "", "", "", "", "Aberto"],
        ["HML-RET-04", "", "", "", "", "Aberto"],
    ], [1.8, 3.0, 5.0, 2.7, 2.0, 1.5])

    add_heading(doc, "7. Resultado da rodada", 2)
    add_table(doc, ["Campo", "Preenchimento"], [
        ["Nome de quem testou", ""],
        ["Data do teste", ""],
        ["Resultado geral", "Aprovado / Aprovado com ressalva / Reprovado"],
        ["Pode encerrar a Etapa 7?", "Sim / Não"],
        ["Resumo", ""],
    ], [5.3, 11.7])
    add_paragraph(doc, "Atenção: esta rodada não substitui a matriz completa. Ela confirma os ajustes da validação anterior. Se algum item reprovar, mantenha o ID e a evidência para correção rastreável.", indent=True)

    add_heading(doc, "Histórico de versões", 2)
    add_table(doc, ["Versão", "Descrição"], [["v1.5", "Roteiro dirigido para revalidar as correções publicadas após o retorno do Tarcísio. Texto revisado em UTF-8, com prints e teste real do QR Code em tela e papel."], ["v1.4", "Matriz anterior de validação final da homologação."]], [2.0, 14.8])

    doc.save(REPO_OUT)
    shutil.copyfile(REPO_OUT, DOWNLOAD_OUT)
    WHATSAPP_OUT.write_text(
        "Pessoal, segue o roteiro v1.5 para revalidar os ajustes do Atenza FieldOps em homologação.\n\n"
        "Por favor, preencham o próprio DOCX com o status, observações e prints. A rodada foca dashboard, ordenação dos contratos, assinatura da proposta, equipe/veículo na agenda, certificados por TAG e QR Code em tela e impresso.\n\n"
        "Não enviem ocorrências soltas no grupo; ao finalizar, devolvam o DOCX preenchido para a Atenza.",
        encoding="utf-8",
    )
    print(REPO_OUT)
    print(DOWNLOAD_OUT)
    print(WHATSAPP_OUT)


if __name__ == "__main__":
    build()
