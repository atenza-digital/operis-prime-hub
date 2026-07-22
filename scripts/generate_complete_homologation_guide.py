from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "cliente" / "homologacao_roteiros"
OUT = OUT_DIR / "Roteiro_Validacao_Completo_Atenza_FieldOps_Ciperprag_v1.6.docx"
IMG_DIR = ROOT / "docs" / "evidencias" / "auditoria-uiux-local"
PRINT_DIR = ROOT / "docs" / "evidencias" / "etapa7_homologacao" / "prints_visuais"
QA_DIR = ROOT / "docs" / "cliente" / "homologacao_roteiros" / "qa_roteiro_completo_v1.6"

GREEN = "087F5B"
NAVY = "142033"
LIGHT = "EAF5F1"
GRAY = "5B6573"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, color=NAVY, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, text in enumerate(headers):
        set_cell_text(header.cells[index], text, bold=True, color="FFFFFF", size=8.5)
        shade(header.cells[index], GREEN)
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            set_cell_text(cells[index], text, size=8.2)
            if len(table.rows) % 2 == 0:
                shade(cells[index], "F7FAF9")
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(GREEN if level == 1 else NAVY)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_step(doc, number, action, expected):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(f"{number}. {action} ").bold = True
    p.add_run(f"Esperado: {expected}")


def add_callout(doc, title, text, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_print(doc, path, caption):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.9))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(8)
    cap.runs[0].font.color.rgb = RGBColor.from_string(GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    for name, size, color in (("Heading 1", 15, GREEN), ("Heading 2", 11.5, NAVY), ("Heading 3", 10, GREEN)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "Atenza FieldOps | Homologação | Roteiro v1.6"
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(GRAY)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Roteiro Completo de Validação")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Atenza FieldOps | Homologação Ciperprag | v1.6")
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(GREEN)
    add_callout(doc, "Objetivo da rodada", "Validar a jornada completa, desde Administração e Comercial até Operacional, Certificados, Relatórios, Medição e acompanhamento no ERP, registrando evidências suficientes para aprovar ou corrigir o produto.")
    add_body(doc, "Este roteiro deve ser executado exclusivamente em homologação. Os dados podem ser de teste e não representam produção. Não registrar senhas em prints, mensagens ou planilhas.")
    add_table(doc, ["Informação", "Valor"], [
        ["URL", "https://fieldops-homologacao.atenza.digital/login"],
        ["Ambiente", "Homologação"],
        ["Versão", "0.6.3"],
        ["Tenant", "Ciperprag"],
        ["Responsável pelo registro", "Tarcísio / equipe Ciperprag"],
    ], [1.4, 5.8])
    add_callout(doc, "Como as senhas serão entregues", "As contas abaixo são contas temporárias. A senha deve ser fornecida separadamente pela Atenza no início da rodada ou regenerada pelo administrador de homologação. O DOCX não grava senhas para evitar vazamento e reutilização indevida.", "FFF7E0")

    add_heading(doc, "1. Contas e perfis de teste")
    add_body(doc, "Use uma conta por área. Se uma conta precisar acumular permissões para executar um cenário de ponta a ponta, registre essa exceção na matriz de aceite.")
    add_table(doc, ["Área", "Usuário", "Perfil", "Senha"], [
        ["Administração", "administrador local / conta admin", "administrativo", "Temporária, fornecida separadamente"],
        ["Comercial", "homolog.comercial@atenza.digital", "comercial", "Temporária, fornecida separadamente"],
        ["Operacional", "homolog.operacao@atenza.digital", "operação + administrativo", "Temporária, fornecida separadamente"],
        ["Qualidade", "homolog.qualidade@atenza.digital", "responsável técnico", "Temporária, fornecida separadamente"],
        ["Medição", "homolog.medicao@atenza.digital", "financeiro", "Temporária, fornecida separadamente"],
    ], [1.2, 2.5, 1.7, 1.8])
    add_bullet(doc, "Troque a senha temporária no primeiro acesso quando o sistema solicitar.")
    add_bullet(doc, "Não use a conta administrativa para validar permissões de perfis menores.")
    add_bullet(doc, "Se a senha for perdida, solicitar regeneração ao administrador da homologação; não criar senha em documento compartilhado.")

    add_heading(doc, "2. Como registrar o resultado")
    add_table(doc, ["Status", "Quando usar"], [
        ["Aprovado", "O comportamento corresponde ao esperado e não houve dúvida relevante."],
        ["Aprovado com observação", "Funciona, mas há melhoria visual ou de usabilidade sem bloquear o fluxo."],
        ["Reprovado", "Impede continuar, grava dado incorreto, mistura tenants ou gera documento inválido."],
        ["Não testado", "Não foi possível executar; registrar o motivo e o passo pendente."],
    ], [2.0, 5.2])
    add_body(doc, "Para cada ocorrência, informe ID do teste, usuário, data/hora, URL, número do documento, passo executado, resultado esperado, resultado encontrado, severidade e anexe um print sem senha.")
    add_callout(doc, "Importante", "Não enviar problemas soltos no grupo de WhatsApp. Preencha este documento ou a matriz de divergências e envie o arquivo consolidado ao final da rodada.", "FFF7E0")

    add_heading(doc, "3. Administração e configuração")
    add_body(doc, "Objetivo: confirmar que o administrador consegue controlar usuários, papéis, permissões, identidade documental, numeração e auditoria sem precisar de suporte técnico.")
    for i, (action, expected) in enumerate([
        ("Entrar com o usuário administrador e confirmar o badge Homologação.", "O ambiente fica claramente identificado como teste."),
        ("Abrir Administração > Usuários e perfis.", "A lista de usuários do tenant aparece sem dados de outro tenant."),
        ("Criar ou editar um usuário de teste e associar um papel.", "O papel é salvo e a permissão é refletida no menu/dashboard."),
        ("Abrir a matriz de permissões e conferir Comercial, Operacional e Financeiro.", "O administrador consegue permitir ou bloquear ações por papel."),
        ("Abrir Configurações > identidade e documentos.", "Logo, ícone, assinatura, cor documental e textos são parametrizáveis por tenant."),
        ("Alterar o último número de uma série em ambiente de teste.", "A próxima numeração segue a sequência configurada sem duplicidade."),
        ("Abrir Eventos de auditoria.", "Alterações de usuários, permissões, documentos e configurações ficam rastreáveis."),
    ], 1):
        add_step(doc, i, action, expected)
    add_print(doc, IMG_DIR / "usu-rios.png", "Print de referência: Usuários e perfis.")
    add_print(doc, PRINT_DIR / "configuracoes-assets-tenant.png", "Print de referência: identidade visual e assets do tenant.")

    add_heading(doc, "4. Comercial: cliente, catálogo, proposta, minuta e contrato")
    add_body(doc, "O fluxo comercial deve começar pelo cadastro ou revisão do cliente e do catálogo. A proposta é enviada e, após aprovação, gera contrato/minuta para liberar a operação.")
    for i, (action, expected) in enumerate([
        ("Abrir Comercial > Clientes e localizar o cliente de teste.", "Razão social, CNPJ, endereço e contatos aparecem corretos."),
        ("Abrir Comercial > Serviços e conferir produtos/serviços.", "Descrição, unidade, recorrência, permite certificado e regras operacionais vêm do catálogo."),
        ("Criar uma proposta com cliente, itens, quantidades, vigência e condições.", "A proposta salva como rascunho sem perder dados."),
        ("Visualizar e imprimir a proposta.", "Layout institucional Ciperprag, Montserrat, acentuação, tabelas, cabeçalho, rodapé e paginação corretos."),
        ("Enviar/aprovar a proposta e gerar a minuta/contrato.", "O contrato é derivado da proposta, sem redigitação desnecessária."),
        ("Conferir cláusulas, periodicidade, reajuste, rescisão, condições comerciais e assinaturas.", "Campos parametrizáveis aparecem coerentes com o tenant e o cliente."),
        ("Visualizar/imprimir contrato e minuta.", "Documentos não se sobrepõem, mantêm a ordem das seções e deixam assinatura legível."),
        ("Usar o caminho Contrato do cliente com um arquivo de referência.", "O sistema aceita o modelo do cliente e integra os itens ao operacional."),
    ], 1):
        add_step(doc, i, action, expected)
    add_print(doc, IMG_DIR / "contratos.png", "Print de referência: Contratos e propostas.")
    add_print(doc, PRINT_DIR / "dashboard-checagem-visual.png", "Print de referência: atalhos e fluxo recomendado.")

    add_heading(doc, "5. Operacional: agenda, equipes e OS")
    add_body(doc, "O Operacional não deve expor valores comerciais ou negociação. Ele recebe contratos vigentes e saldo operacional, organiza agenda/equipe/veículo e conduz a execução de campo.")
    for i, (action, expected) in enumerate([
        ("Abrir Agendamentos e filtrar mês, semana, ano e cliente.", "Calendário e filtros mostram o período escolhido, em formato brasileiro."),
        ("Criar agendamento a partir de contrato vigente.", "Somente itens com saldo operacional disponível aparecem."),
        ("Definir data, local, técnicos, veículo e tags/equipamentos.", "A equipe e os recursos ficam visíveis no detalhe do agendamento."),
        ("Clicar no evento do calendário.", "Abre detalhe contextual sem parecer que uma tela foi empilhada sobre outra."),
        ("Gerar OS e imprimir a via de campo.", "OS usa numeração automática, dados dinâmicos e layout Ciperprag aprovado."),
        ("Abrir Ordens de serviço e conferir cliente, serviço, local, equipe e tag.", "A OS corresponde exatamente ao agendamento e ao contrato."),
        ("Encerrar a OS com data, quantidade, tag, checklist e até três fotos.", "A OS encerra sem erro de API e as fotos ficam vinculadas ao tenant/OS."),
        ("Registrar não execução quando aplicável.", "O motivo fica registrado e não gera baixa indevida."),
    ], 1):
        add_step(doc, i, action, expected)
    add_print(doc, IMG_DIR / "agendamentos.png", "Print de referência: Agendamentos.")
    add_print(doc, IMG_DIR / "ordens.png", "Print de referência: Ordens de serviço.")

    add_heading(doc, "6. Qualidade: certificado, QR Code, histórico e relatório")
    add_body(doc, "O certificado deve ser derivado da mesma OS, do mesmo tenant e do mesmo snapshot. Não pode misturar cliente, serviço, título, produtos, fotos ou responsável de registros diferentes.")
    for i, (action, expected) in enumerate([
        ("Abrir Certificados e histórico.", "Serviços encerrados aparecem; serviços sem certificado também permanecem no histórico."),
        ("Gerar certificado de uma OS cujo serviço permite emissão.", "Cliente, CNPJ, serviço, OS, execução, tag, local, fotos e responsável são coerentes."),
        ("Conferir o quadro de autenticidade.", "Número do certificado, OS, data, QR Code, código curto e fingerprint SHA-256 aparecem legíveis."),
        ("Ler o QR Code em tela e em papel A4.", "A rota pública abre exatamente o certificado correspondente."),
        ("Testar certificado com zero, uma e até três fotos.", "Fotos não deformam, não criam espaços vazios indevidos e não ultrapassam o limite."),
        ("Abrir Relatórios técnicos e gerar um relatório.", "O relatório usa dados da OS e mantém cabeçalho/rodapé e acentuação."),
    ], 1):
        add_step(doc, i, action, expected)
    add_print(doc, IMG_DIR / "certificados.png", "Print de referência: Certificados e histórico.")
    add_print(doc, PRINT_DIR / "agenda-visao-mensal.png", "Print de referência: navegação visual em homologação.")
    add_callout(doc, "Teste antifraude", "Use pelo menos dois aparelhos para ler QR Codes de dois certificados diferentes. Registre o número retornado e confirme que não houve troca entre documentos.")

    add_heading(doc, "7. Financeiro operacional: medição e acompanhamento até o ERP")
    add_body(doc, "Este módulo não substitui contas a pagar/receber. Ele consolida serviços executados, gera medição e acompanha NF, pagamento e necessidade de cobrança até a baixa manual no ERP.")
    for i, (action, expected) in enumerate([
        ("Abrir Medição e selecionar cliente e intervalo de datas.", "A tela mostra somente OS encerradas, elegíveis e ainda não medidas."),
        ("Gerar a medição.", "Itens, quantidades, unidades, valores e total correspondem às OS e ao contrato."),
        ("Visualizar/imprimir o PDF.", "PDF A4 retrato, Montserrat incorporada, sem cortes, com logo e dados dinâmicos do tenant."),
        ("Alterar para Aguardando NF e registrar número/data de envio.", "O acompanhamento fica salvo e auditável."),
        ("Alterar para NF enviada, Aguardando pagamento e Pago no ERP.", "O status acompanha a situação sem criar contas financeiras paralelas."),
        ("Testar uma medição com vários itens e uma com nomes extensos.", "Paginação, totais, assinaturas e rastreabilidade permanecem íntegros."),
    ], 1):
        add_step(doc, i, action, expected)
    add_print(doc, IMG_DIR / "medi-o.png", "Print de referência: Medição.")
    add_print(doc, PRINT_DIR / "medicao-checagem-visual.png", "Print de referência: medição em validação visual.")

    add_heading(doc, "8. Recorrência e continuidade do fluxo")
    for i, (action, expected) in enumerate([
        ("Encerrar uma OS de serviço recorrente.", "O sistema calcula uma sugestão de próxima data."),
        ("Abrir Dashboard ou Agendamentos e revisar a sugestão.", "Cliente, serviço, intervalo e data sugerida ficam claros."),
        ("Confirmar a sugestão.", "Um novo agendamento entra na agenda sem refazer o cadastro."),
        ("Dispensar uma sugestão.", "Nenhum agendamento indevido é criado e a ação fica registrada quando aplicável."),
    ], 1):
        add_step(doc, i, action, expected)

    add_heading(doc, "9. Isolamento SaaS e identidade documental")
    add_body(doc, "Esta rodada deve confirmar que Ciperprag não é uma marca fixa do produto. A tela de login é Atenza FieldOps; após o login, sidebar e documentos usam os assets do tenant. O teste deve ser executado também com empresa demonstração e tenant sem logo, quando as contas estiverem disponíveis.")
    add_table(doc, ["Cenário", "Conferir"], [
        ["Ciperprag", "Logo, CNPJ, contatos, cor documental e assinatura da Ciperprag apenas em documentos do tenant."],
        ["Empresa demonstração", "Logo/dados fictícios, sem reaproveitar CNPJ, telefone ou e-mail Ciperprag."],
        ["Tenant sem logo", "Layout neutro, sem retângulo quebrado, sem fallback fixo da Ciperprag."],
    ], [2.0, 5.2])
    add_print(doc, PRINT_DIR / "login-saas-sem-ciperprag.png", "Print de referência: login neutro da plataforma SaaS.")
    add_print(doc, PRINT_DIR / "dashboard-sidebar-logo-tenant.png", "Print de referência: logo do tenant após autenticação.")

    add_heading(doc, "10. Matriz consolidada de aceite")
    rows = []
    matrix = [
        ("ADM-01", "Login e ambiente", "Badge Homologação e versão visíveis", ""),
        ("ADM-02", "Usuários e perfis", "Papel e permissões respeitados", ""),
        ("ADM-03", "Assets e numeração", "Configuração salva e aplicada", ""),
        ("COM-01", "Cliente e catálogo", "Dados usados nos documentos", ""),
        ("COM-02", "Proposta", "Criar, salvar, imprimir e aprovar", ""),
        ("COM-03", "Minuta/contrato", "Gerar após aprovação ou importar cliente", ""),
        ("OP-01", "Agendamento", "Contrato vigente, saldo, equipe e veículo", ""),
        ("OP-02", "OS", "Gerar, imprimir, editar e conferir", ""),
        ("OP-03", "Encerramento", "Data, quantidade, tag, checklist e até 3 fotos", ""),
        ("QL-01", "Certificado", "Dados da mesma OS, logo e Montserrat", ""),
        ("QL-02", "QR/histórico", "Validação pública e serviços sem certificado", ""),
        ("QL-03", "Relatório", "Dados técnicos e PDF íntegro", ""),
        ("FIN-01", "Medição", "OS elegíveis por período e baixa contratual", ""),
        ("FIN-02", "Status financeiro", "NF, pagamento, cobrança e ERP", ""),
        ("OP-04", "Recorrência", "Sugestão confirmada vira agenda", ""),
        ("SAAS-01", "Isolamento", "Sem vazamento entre tenants", ""),
        ("DOC-01", "Documentos", "Paginação, acentos, logo, rodapé e assinaturas", ""),
        ("R2-01", "Anexos", "Visualização/download e SHA-256 íntegro", ""),
    ]
    for test_id, item, expected, _ in matrix:
        rows.append([test_id, item, expected, "Pendente", ""])
    add_table(doc, ["ID", "Item", "Resultado esperado", "Status", "Observação/evidência"], rows, [0.65, 1.35, 3.05, 0.85, 1.3])

    add_heading(doc, "11. Registro de problemas")
    add_table(doc, ["ID", "Perfil", "Passo", "Problema", "Severidade", "Status"], [
        ["HML-001", "", "", "", "Alta/Média/Baixa", "Aberto"],
        ["HML-002", "", "", "", "Alta/Média/Baixa", "Aberto"],
        ["HML-003", "", "", "", "Alta/Média/Baixa", "Aberto"],
    ], [0.8, 1.0, 0.8, 2.7, 1.1, 0.8])
    add_body(doc, "Um problema é Alta quando impede o fluxo, mistura dados, perde evidência ou gera documento incorreto; Média quando há contorno seguro, mas a tela confunde; Baixa quando é texto, acentuação, alinhamento ou melhoria de conveniência.")
    add_callout(doc, "Mensagem para o grupo", "A equipe pode usar a homologação para validar o fluxo completo do Atenza FieldOps. Por favor, registrem cada resultado no roteiro e enviem o DOCX preenchido ao final. Não enviem ocorrências soltas no grupo de WhatsApp, pois o documento facilita o mapeamento e a correção.")

    add_heading(doc, "12. Critério de encerramento da rodada")
    add_bullet(doc, "Todos os IDs da matriz foram classificados como Aprovado, Aprovado com observação, Reprovado ou Não testado.")
    add_bullet(doc, "Toda reprovação possui print, número do documento ou URL e descrição reproduzível.")
    add_bullet(doc, "Documentos foram conferidos visualmente, incluindo OS, proposta, minuta, contrato, certificado, relatório e medição.")
    add_bullet(doc, "QR Codes foram lidos em tela e em papel quando possível.")
    add_bullet(doc, "A equipe enviou o DOCX preenchido e os anexos de evidência para consolidação pela Atenza.")

    doc.core_properties.title = "Roteiro Completo de Validação Atenza FieldOps - Ciperprag"
    doc.core_properties.subject = "Homologação funcional, visual e SaaS"
    doc.core_properties.author = "Atenza"
    doc.core_properties.comments = "Versão 1.6 - homologação"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
