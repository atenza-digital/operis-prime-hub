# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = (
    Path("C:/Projetos")
    / "Documentações"
    / "02_Templates"
    / "01_Operacionais"
    / "Template_Checklist_Prontidao_Producao_Atenza.docx"
)
OUT_DIR = ROOT / "docs" / "cliente" / "relatorios_homologacao"
QA_DIR = ROOT / "docs" / "evidencias" / "etapa7_homologacao" / "relatorio_prontidao_p0"
DOCX_OUT = OUT_DIR / "Relatorio_Prontidao_P0_Atenza_FieldOps_v1.0.docx"
QA_OUT = QA_DIR / "QA_RELATORIO_PRONTIDAO_P0.md"

PROJECT = "Atenza FieldOps"
TENANT = "Ciperprag"
URL = "https://fieldops-homologacao.atenza.digital/login"
VERSION = "0.6.3"
BRANCH = "homologacao/p0-relatorios-tecnicos"
PR = "#4 - Integrar fluxo P0 para homologação"
COMMIT = "4098773"
RESPONSIBLE = "Ewerton Gomes Almeida"
DATE_BR = datetime.now().strftime("%d/%m/%Y")


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_font(run, *, bold: bool = False, size: float = 10.0) -> None:
    run.font.name = "Nortica Typeface Bold" if bold else "Nortica Typeface"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.size = Pt(size)
    run.bold = bold


def paragraph(
    doc: Document,
    text: str = "",
    *,
    justify: bool = True,
    indent: bool = True,
    bold: bool = False,
    size: float = 10.0,
):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0.7) if indent else Cm(0)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    set_font(run, bold=bold, size=size)
    return p


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Neue Power Ultra" if level == 1 else "Nortica Typeface Bold"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), run.font.name)
    run.font.size = Pt(18 if level == 1 else 13)
    run.bold = True
    return p


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 8.8) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    set_font(run, bold=bold, size=size)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        shade(cell, "F2F2ED")
    for row in rows:
        cells = tbl.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths_cm:
        for row in tbl.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return tbl


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"- {text}")
    set_font(run, size=9.5)


def numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(4)
    count = getattr(numbered, "_count", 0) + 1
    numbered._count = count
    run = p.add_run(f"{count}. {text}")
    set_font(run, size=9.5)


def reset_numbered() -> None:
    numbered._count = 0


def normalize_template_text_fonts(doc: Document) -> None:
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                for run in p.runs:
                    set_font(run, size=8)


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document(str(TEMPLATE))
    clear_body(doc)
    normalize_template_text_fonts(doc)
    doc.core_properties.title = "Relatório de Prontidão P0 - Atenza FieldOps"
    doc.core_properties.author = "Atenza"
    doc.core_properties.subject = "Prontidão de homologação do fluxo P0 Ciperprag"
    doc.core_properties.comments = "Documento de acompanhamento de homologação, sem credenciais ou segredos."

    paragraph(doc, "DOCUMENTO INSTITUCIONAL", justify=False, indent=False, bold=True, size=9)
    heading(doc, "Relatório de Prontidão P0", 1)
    paragraph(
        doc,
        "Atenza FieldOps - consolidação executiva da prontidão de homologação assistida do fluxo P0 da Ciperprag.",
        justify=False,
        indent=False,
    )

    heading(doc, "Ficha executiva", 2)
    table(
        doc,
        ["Campo", "Informação"],
        [
            ["Projeto", PROJECT],
            ["Tenant de homologação", TENANT],
            ["Ambiente", "Homologação"],
            ["URL oficial", URL],
            ["Versão", VERSION],
            ["Branch", BRANCH],
            ["PR", PR],
            ["Commit-base deste relatório", COMMIT],
            ["Responsável Atenza", RESPONSIBLE],
            ["Data do relatório", DATE_BR],
            ["Status executivo", "Pronto para homologação assistida, ainda não liberado para produção SaaS."],
        ],
        [4.7, 11.0],
    )

    heading(doc, "1. Decisão recomendada", 2)
    paragraph(
        doc,
        "A recomendação é manter a Etapa 7 aberta para a validação assistida com Tarcísio e equipe, usando o roteiro final já gerado. O P0 está tecnicamente organizado para teste ponta a ponta em homologação, mas a liberação para produção deve aguardar a devolutiva humana, a triagem das divergências e a etapa de hardening SaaS.",
    )
    bullet(doc, "Liberar uso em homologação: recomendado.")
    bullet(doc, "Liberar produção SaaS: não recomendado neste momento.")
    bullet(doc, "Avançar em paralelo: permitido apenas em itens que não alterem o fluxo em teste.")

    heading(doc, "2. Prontidão por frente do P0", 2)
    table(
        doc,
        ["Frente P0", "Situação", "Critério principal", "Evidência"],
        [
            [
                "1. Propostas",
                "Pronta para homologação",
                "Proposta com layout aprovado, ordem de seções corrigida, Montserrat incorporada e fluxo para minuta/contrato.",
                "docs/evidencias/qa_fluxo_visual/proposta-ciperprag.pdf",
            ],
            [
                "2. Contratos e minutas",
                "Pronto para homologação",
                "Fluxo proposta -> minuta -> contrato final -> contrato operacional validado tecnicamente.",
                "docs/evidencias/qa_fluxo_visual/contrato-ciperprag.pdf",
            ],
            [
                "3. Agendamentos",
                "Pronto para homologação",
                "Agenda usa contratos vigentes com saldo e permite equipe, veículo, local e detalhes do serviço.",
                "docs/evidencias/etapa7_homologacao/execucao-tecnica-e2e.md",
            ],
            [
                "4. Ordens de Serviço",
                "Pronta para homologação",
                "OS impressa pelo gerador centralizado, com tag/equipamento, checklist e fechamento com evidências.",
                "docs/evidencias/qa_fluxo_visual/os-atual.pdf",
            ],
            [
                "5. Certificados",
                "Pronto para homologação",
                "Certificado em Montserrat, QR Code, hash SHA-256, fotos dinâmicas, blocos condicionais e validação pública.",
                "docs/cliente/certificados_montserrat/certificado-ciperprag-amostra-final.pdf",
            ],
            [
                "6. Relatórios técnicos",
                "Pronto para homologação",
                "Relatório a partir da OS, sem valores comerciais, com checklist, fotos, produtos, EPIs e normas.",
                "docs/cliente/relatorios_tecnicos/relatorio-tecnico-ciperprag-amostra.pdf",
            ],
            [
                "7. Medições e acompanhamento",
                "Pronta para homologação",
                "Medição em A4 retrato, totalização em centavos, status de NF/pagamento/ERP e Montserrat incorporada.",
                "docs/evidencias/etapa7_homologacao/medicoes/med-validacao-2026-a4-retrato.pdf",
            ],
        ],
        [3.1, 3.2, 5.3, 4.1],
    )

    heading(doc, "3. Evidências técnicas consolidadas", 2)
    table(
        doc,
        ["Evidência", "Resumo"],
        [
            [
                "Roteiro final Tarcísio v1.1",
                "Documento formal para validação assistida por perfil, com checklist e registro de divergências.",
            ],
            [
                "Auditoria de fontes documentais",
                "Proposta, contrato, OS, certificado, relatório técnico e medição validados com Montserrat incorporada.",
            ],
            [
                "Smoke E2E local",
                "Fluxo proposta aprovada, minuta, contrato, agenda, OS, certificado, medição e recorrência aprovado tecnicamente.",
            ],
            [
                "Smoke VPS de homologação",
                "Health, login, API protegida, bootstrap e certificado público já foram cobertos em validações anteriores.",
            ],
            [
                "Saneamento de homologação",
                "Dados legados com encoding e resíduos funcionais antigos foram tratados no tenant Ciperprag.",
            ],
        ],
        [5.0, 10.7],
    )

    heading(doc, "4. Itens que ainda exigem validação humana", 2)
    reset_numbered()
    numbered(doc, "Executar o fluxo completo com usuário humano: proposta, minuta, contrato final, agenda, OS, encerramento, certificado, medição e recorrência.")
    numbered(doc, "Validar leitura real de QR Code em tela e impresso com celular, porque esse ponto depende de experiência física.")
    numbered(doc, "Conferir se os nomes, rótulos, acentuação, datas, horas e valores fazem sentido para a rotina da Ciperprag.")
    numbered(doc, "Registrar divergências no DOCX do roteiro, com print, número do documento e severidade.")
    numbered(doc, "Confirmar se o fluxo está simples o suficiente para adoção pelos usuários sem orientação constante.")

    heading(doc, "5. Bloqueios para produção", 2)
    table(
        doc,
        ["Bloqueio", "Motivo", "Etapa"],
        [
            [
                "Feedback humano pendente",
                "A equipe ainda precisa executar a rodada assistida e devolver o roteiro preenchido.",
                "Etapa 7",
            ],
            [
                "Hardening SaaS",
                "Produção exige isolamento final, storage, políticas de segurança, observabilidade e rollback formal.",
                "Etapa 8",
            ],
            [
                "PDF server-side final",
                "Os documentos já têm padrão visual, mas produção deve persistir PDF binário imutável com snapshot e hash.",
                "Etapa 8",
            ],
            [
                "Governança multi-tenant",
                "Ainda faltam validações finais de tenant genérico, tenant sem logo e parametrização documental completa.",
                "Etapa 8",
            ],
            [
                "PR com política de review",
                "Merge na main depende de aprovação conforme política do repositório.",
                "Governança Git",
            ],
        ],
        [4.2, 8.2, 3.3],
    )

    heading(doc, "6. Backlog remanescente consolidado", 2)
    paragraph(
        doc,
        "O backlog permanece centralizado no roadmap por etapas. A contagem atual é de 46 itens remanescentes, sendo 8 na Etapa 7 e 38 na Etapa 8. Não há itens fora de etapa.",
    )
    table(
        doc,
        ["Grupo", "Quantidade", "Observação"],
        [
            ["Etapa 7 - QA e homologação guiada", "8", "Itens de validação manual, multi-tenant, numeração, anexos, UX e QR Code físico."],
            ["Etapa 8 - Produção e hardening SaaS", "38", "Itens de produção, governança, storage, R2, PDF server-side, segurança, permissões e UX avançada."],
            ["Itens fora de etapa", "0", "Todos os itens conhecidos estão alocados no roadmap."],
        ],
        [5.2, 2.6, 7.9],
    )

    heading(doc, "7. Próxima ação sugerida", 2)
    paragraph(
        doc,
        "A próxima ação sugerida é aguardar a devolução do roteiro preenchido pelo Tarcísio e, em paralelo, preparar a matriz de triagem das divergências para classificar cada retorno como correção P0 obrigatória, ajuste de UX da Etapa 8 ou melhoria futura parametrizável. Isso evita retrabalho e mantém o P0 fechado por evidência, não por percepção.",
    )

    heading(doc, "Histórico de versões", 2)
    table(
        doc,
        ["Versão", "Data", "Responsável", "Descrição"],
        [
            ["v1.0", DATE_BR, RESPONSIBLE, "Primeira consolidação executiva de prontidão P0 para homologação assistida."],
        ],
        [2.0, 2.8, 4.2, 7.0],
    )

    doc.save(DOCX_OUT)

    QA_OUT.write_text(
        "\n".join(
            [
                "# QA - Relatório de Prontidão P0",
                "",
                f"- Documento: `{DOCX_OUT}`",
                "- PDF: pendente de renderização pós-geração.",
                "- Base visual: `Template_Checklist_Prontidao_Producao_Atenza.docx`.",
                "- Status inicial: DOCX gerado, aguardando exportação PDF e inspeção PNG.",
                "- Checklist: acentuação, layout, rodapé, tabelas, ausência de credenciais e clareza executiva.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    print(DOCX_OUT)
    print(QA_OUT)


if __name__ == "__main__":
    build()
