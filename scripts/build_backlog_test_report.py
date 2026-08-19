from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "cliente" / "relatorios_homologacao" / "Relatorio_Testes_Estoque_PDF_Locais_Consumo_Atenza_FieldOps_2026-08-13.docx"
ACCENT = "087F5B"
INK = "162235"
MUTED = "52627A"
LIGHT = "EAF4F0"


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D8E2E8", size="6"):
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_font(run, size=10.5, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_font(run, **kwargs)
    return run


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    paragraph.paragraph_format.space_after = Pt(6)
    add_text(paragraph, text, size=15 if level == 1 else 11.5, bold=True, color=ACCENT if level == 1 else INK)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        add_text(paragraph, bold_prefix, bold=True)
        add_text(paragraph, text[len(bold_prefix):])
    else:
        add_text(paragraph, text)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.12)
        add_text(paragraph, item, size=10)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, ACCENT)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            cell.width = Inches(widths[index])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_text(paragraph, header, size=9, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell = cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cell.width = Inches(widths[index])
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_text(paragraph, str(value), size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_font(run, size=8, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(header, "ATENZA FIELDOPS  |  RELATÓRIO DE HOMOLOGAÇÃO", size=8, bold=True, color=ACCENT)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "Atenza FieldOps  |  Homologação  |  ", size=8, color=MUTED)
    add_page_number(footer)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    add_text(title, "Relatório de testes", size=22, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    add_text(subtitle, "Estoque, importação de PDF, locais de execução, consumo e PDFs server-side", size=11.5, color=MUTED)

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(callout, color="B7DACA", size="8")
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    p = cell.paragraphs[0]
    add_text(p, "Resultado geral: APROVADO TECNICAMENTE EM DESENVOLVIMENTO", size=10.5, bold=True, color=ACCENT)
    p = cell.add_paragraph()
    add_text(p, "A execução no container de homologação e a validação visual externa permanecem como aceite final via CI/CD.", size=9.5, color=INK)

    add_heading(doc, "1. Escopo desta rodada", 1)
    add_body(doc, "Implementar as pendências recebidas da validação da Ciperprag relacionadas a estoque, importação de propostas por PDF, locais cadastrados, histórico de consumo e geração server-side de documentos, sem usar dados locais como fonte funcional.")
    add_table(doc, ["Frente", "Resultado tecnico", "Status"], [
        ("Estoque", "Saldo, movimentos auditados e baixa transacional por OS/serviço.", "Aprovado"),
        ("PDF de proposta", "Extração local determinística, cobertura e original preservado.", "Aprovado"),
        ("Locais", "Seleção cadastrada no agendamento e propagação para OS/recorrência.", "Aprovado"),
        ("Consumo", "Histórico e resumo filtrável por período, produto e OS.", "Aprovado"),
        ("Documentos", "Renderização Chromium e anexo PDF imutável com hash.", "Aprovado tecnicamente"),
    ], widths=[1.15, 4.3, 1.0])

    add_heading(doc, "2. Testes automatizados e técnicos", 1)
    add_table(doc, ["Verificação", "Resultado", "Evidência"], [
        ("TypeScript", "Aprovado", "npx tsc --noEmit"),
        ("ESLint", "Aprovado; 1 warning preexistente", "npm run lint"),
        ("Vitest", "41 testes aprovados", "8 arquivos de teste"),
        ("Build web", "Aprovado", "npm run build"),
        ("PDF determinístico", "4 páginas, 2 tabelas, 10 linhas", "test:proposal-pdf"),
        ("PDF server-side", "Assinatura %PDF-; 30.881 bytes", "test:server-pdf"),
        ("Smoke estoque", "Entrada 12, saída 3, saldo 9; limpeza", "homologation:stock"),
    ], widths=[1.5, 2.2, 2.75])

    add_heading(doc, "3. Cenários cobertos", 1)
    add_bullets(doc, [
        "Produto de estoque por tenant com unidade, saldo inicial, estoque mínimo e ativo/inativo.",
        "Movimentos de entrada, saída, ajuste, devolução e perda registrados com saldo anterior/posterior, usuário, OS e serviço.",
        "Baixa de consumo no encerramento da OS com bloqueio de saldo insuficiente e proteção contra duplicidade da saída da mesma OS.",
        "PDF de referência analisado primeiro localmente; a chamada de IA continua produzindo apenas um rascunho revisável.",
        "Hash SHA-256, conteúdo original, texto extraído, páginas e cobertura persistidos na tabela de importações.",
        "Locais do cliente selecionados por identificador no agendamento, mantidos na OS e na sugestão de recorrência.",
        "Relatório de consumo com filtros de data, produto e OS, além de resumo por tipo de movimento.",
        "PDFs históricos gerados no servidor, gravados como application/pdf, com hash e imutabilidade.",
    ])

    add_heading(doc, "4. Validação pendente em homologação", 1)
    add_table(doc, ["Teste manual", "Como validar", "Resultado esperado"], [
        ("Estoque", "Cadastrar produto, movimentar entrada e conferir saldo.", "Saldo e histórico atualizados sem dados locais."),
        ("Baixa por OS", "Encerrar OS com produto previsto e quantidade usada.", "Saída vinculada a OS/serviço; saldo reduzido uma vez."),
        ("Importar proposta", "Anexar PDF real e revisar o rascunho antes de salvar.", "Tabelas/campos reconhecidos; original continua recuperável."),
        ("Local", "Cadastrar dois locais e selecionar um no agendamento.", "OS e recorrência preservam o local escolhido."),
        ("Consumo", "Filtrar por período e por número de OS.", "Movimentos e resumo refletem o filtro."),
        ("Documentos", "Gerar OS, proposta, certificado, contrato e medição.", "PDF abre, mantém layout aprovado e possui hash."),
    ], widths=[1.25, 3.0, 2.2])
    add_body(doc, "URL de homologação: https://fieldops-homologacao.atenza.digital/login")
    add_body(doc, "Usuário de teste: homolog.operacao@atenza.digital. A senha deve ser informada pelo canal seguro já utilizado pela equipe e não deve ser registrada neste relatório.")

    add_heading(doc, "5. Pendências e próxima rodada", 1)
    add_bullets(doc, [
        "Executar o smoke de estoque dentro do container publicado pela pipeline de homologação, depois da revisão e merge do PR.",
        "Conferir visualmente os cinco PDFs server-side contra as referências Ciperprag aprovadas.",
        "Tarcísio/Aline revalidarem estoque, importação de PDF, locais, consumo e documentos, registrando aprovado, observação, reprovado ou não testado.",
        "Manter como backlog posterior: R2 definitivo, templates versionados, histórico de versões, backfill de documentos antigos e hardening de antivírus.",
    ])
    add_heading(doc, "6. Arquivos e evidências", 1)
    add_bullets(doc, [
        "docs/evidencias/etapa9_homologacao/BACKLOG_ESTOQUE_PDF_LOCAIS_CONSUMO_2026-08-13.md",
        "scripts/smoke-stock-homologation.mjs",
        "scripts/test-deterministic-proposal-pdf.mjs",
        "scripts/test-server-pdf.mjs",
        "server/render-pdf.mjs",
    ])
    add_body(doc, "Nenhuma alteração foi aplicada diretamente na VPS. A publicação deve ocorrer exclusivamente pelo workflow de CI/CD. A renderização visual deste DOCX não foi executada porque o ambiente não possui LibreOffice/soffice instalado; a estrutura do arquivo e o conteúdo foram verificados programaticamente.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
